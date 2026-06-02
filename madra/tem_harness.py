from __future__ import annotations

from collections import Counter
import json
from pathlib import Path
import re
from statistics import mean, median, pstdev
from typing import Any

from .io import read_jsonl
from .paper_writing import audit_manuscript, build_reference_library


FORMAL_RESULTS_MARKER = "[FORMAL_500_RESULTS_TO_BE_UPDATED]"

FIGURE_CAPTIONS = {
    "fig1_madra_architecture": "Fig. 1. Blackboard-mediated MAD-RA architecture for evidence-grounded responsibility attribution.",
    "fig2_deliberation_workflow": "Fig. 2. Structured message passing and shared case blackboard state.",
    "fig3_pilot_protocol": "Fig. 3. Conflict detection and targeted re-deliberation workflow.",
    "fig4_dataset_profile": "Fig. 4. Formal 500-case pilot protocol, baselines, and validation layers.",
    "fig5_results_dashboard": "Fig. 5. Evaluation dashboard for the formal 500-case full MAD-RA pilot.",
    "fig6_management_loop": "Fig. 6. Workflow-level auditability and management decision-support outputs.",
}


def run_tem_harness(
    *,
    dataset_path: str | Path = "outputs/pilot_500/dataset_madra_500.jsonl",
    output_dir: str | Path = "outputs/ieee_tem_harness",
    human_checked_path: str | Path | None = "outputs/pilot_500/human_checked_50.jsonl",
    metrics_path: str | Path | None = "outputs/pilot_500/metrics_500.json",
    calibration_report_path: str | Path | None = "outputs/pilot_500/stage2_qwen_30_calibration_report_clean.json",
    calibration_metrics_path: str | Path | None = "outputs/pilot_500/stage2_qwen_30_calibration_metrics_clean.json",
    token_usage_path: str | Path | None = "outputs/pilot_500/token_usage_500.json",
    runtime_path: str | Path | None = "outputs/pilot_500/runtime_500.json",
    reference_path: str | Path | None = "outputs/pilot_500/MADRA_61refs.ris",
    format_reference_docx: str | Path | None = None,
    write_docx: bool = True,
) -> dict[str, Path]:
    output = Path(output_dir)
    output.mkdir(parents=True, exist_ok=True)

    rows = read_jsonl(dataset_path)
    human_rows = read_jsonl(human_checked_path) if human_checked_path and Path(human_checked_path).exists() else []
    profile = build_dataset_profile(rows, human_checked_rows=human_rows)
    metrics = read_json(metrics_path)
    calibration_report = read_json(calibration_report_path)
    calibration_metrics = read_json(calibration_metrics_path)
    token_usage = read_json(token_usage_path)
    runtime = read_json(runtime_path)
    references = build_reference_library(reference_path=reference_path)

    data_profile_path = output / "dataset_profile_500.json"
    data_profile_path.write_text(json.dumps(profile, ensure_ascii=False, indent=2), encoding="utf-8")
    figure_contract_path = output / "figure_contract.md"
    figure_contract_path.write_text(build_figure_contract(metrics=metrics), encoding="utf-8")
    prompt_path = output / "figure_prompts.md"
    prompt_path.write_text(build_ai_figure_prompts(), encoding="utf-8")
    absorption_path = output / "ppt_absorption_checklist.md"
    absorption_path.write_text(build_ppt_absorption_checklist(), encoding="utf-8")

    figure_dir = output / "figures"
    figure_paths = make_tem_figures(
        output_dir=figure_dir,
        profile=profile,
        metrics=metrics,
        calibration={**calibration_metrics, **calibration_report},
        token_usage=token_usage,
        runtime=runtime,
    )

    manuscript = build_complete_manuscript(
        profile=profile,
        metrics=metrics,
        calibration_report=calibration_report,
        calibration_metrics=calibration_metrics,
        token_usage=token_usage,
        runtime=runtime,
        references=references,
        figure_paths=figure_paths,
    )
    manuscript_path = output / "IEEE_TEM_MADRA_full_draft.md"
    manuscript_path.write_text(manuscript, encoding="utf-8")

    audit = audit_manuscript(manuscript, references)
    audit.update(
        {
            "formal_500_results_present": bool(metrics),
            "formal_results_marker_present": FORMAL_RESULTS_MARKER in manuscript,
            "dataset_cases": profile["num_cases"],
            "figure_count": len(figure_paths),
        }
    )
    audit_path = output / "manuscript_audit.json"
    audit_path.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")

    update_manifest_path = output / "rerun_update_manifest.md"
    update_manifest_path.write_text(build_update_manifest(), encoding="utf-8")

    docx_path = output / "IEEE_TEM_MADRA_full_draft.docx"
    if write_docx:
        write_tem_docx(
            manuscript,
            docx_path,
            figure_paths=figure_paths,
            format_reference_docx=format_reference_docx,
        )

    return {
        "manuscript_path": manuscript_path,
        "audit_path": audit_path,
        "data_profile_path": data_profile_path,
        "figure_contract_path": figure_contract_path,
        "figure_prompts_path": prompt_path,
        "ppt_absorption_checklist_path": absorption_path,
        "update_manifest_path": update_manifest_path,
        "docx_path": docx_path,
    }


def write_tem_docx(
    manuscript: str,
    output_path: str | Path,
    *,
    figure_paths: list[Path],
    format_reference_docx: str | Path | None = None,
) -> Path:
    """Write an AIC/IEEE-TEM-style manuscript DOCX with real tables and embedded figures."""
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION_START
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Inches, Pt, RGBColor
    except Exception as exc:  # pragma: no cover - dependency guard
        raise RuntimeError("python-docx is required to write the formatted MAD-RA manuscript DOCX.") from exc

    def cell_shading(cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_mar = tc_pr.first_child_found_in("w:tcMar")
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for margin_name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
            node = tc_mar.find(qn(f"w:{margin_name}"))
            if node is None:
                node = OxmlElement(f"w:{margin_name}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    def set_table_width(table, width_dxa: int) -> None:
        table.autofit = False
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), str(width_dxa))
        tbl_w.set(qn("w:type"), "dxa")
        tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
        if tbl_ind is None:
            tbl_ind = OxmlElement("w:tblInd")
            tbl_pr.append(tbl_ind)
        tbl_ind.set(qn("w:w"), "0")
        tbl_ind.set(qn("w:type"), "dxa")

    def style_run(run, size: float = 12.0, bold: bool | None = None, italic: bool | None = None) -> None:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic

    def style_paragraph(paragraph, *, size: float = 12.0, bold: bool = False, italic: bool = False) -> None:
        for run in paragraph.runs:
            style_run(run, size=size, bold=bold, italic=italic)

    def add_caption(doc, text: str) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(text)
        style_run(run, size=12.0)

    def add_body_paragraph(doc, text: str) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        style_run(run, size=12.0)

    def add_heading(doc, text: str, level: int) -> None:
        if level <= 1:
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(text)
            style_run(run, size=float(tokens.get("title_size", 16.0)), bold=True)
            return
        p = doc.add_paragraph(style="Heading 1" if level == 2 else "Heading 2")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(8 if level == 2 else 4)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5 if level == 2 else 1.15
        run = p.add_run(text)
        style_run(run, size=12.0, bold=True)

    def parse_table_line(line: str) -> list[str]:
        stripped = line.strip().strip("|")
        return [cell.strip().replace("<br>", "\n") for cell in stripped.split("|")]

    def is_separator_row(cells: list[str]) -> bool:
        return all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)

    def add_markdown_table(doc, rows: list[list[str]], table_idx: int) -> None:
        rows = [row for row in rows if row and not is_separator_row(row)]
        if not rows:
            return
        captions = [
            "Table 1. Dataset profile for the 500-case pilot subset.",
            "Table 2. Candidate-label distribution in the pilot subset.",
            "Table 3. Stage-2 Qwen calibration checks.",
            "Table 4. Baseline and ablation design.",
            "Table 5. Evaluation metrics and managerial interpretation.",
            "Table 6. Runtime, token, and cost diagnostics.",
            "Table 7. Formal 500-case pilot metrics.",
            "Table 8. Figure files and editable export formats.",
        ]
        add_caption(doc, captions[table_idx - 1] if table_idx <= len(captions) else f"Table {table_idx}. Manuscript table.")
        col_count = max(len(row) for row in rows)
        table = doc.add_table(rows=len(rows), cols=col_count)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_width(table, 9360)
        for r_idx, row in enumerate(rows):
            for c_idx in range(col_count):
                cell = table.cell(r_idx, c_idx)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                text = row[c_idx] if c_idx < len(row) else ""
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 18 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(text)
                style_run(run, size=9.5, bold=(r_idx == 0))
                if r_idx == 0:
                    cell_shading(cell, "E8EEF5")
        doc.add_paragraph().paragraph_format.space_after = Pt(3)

    def add_figures(doc) -> None:
        add_heading(doc, "Embedded manuscript figures", 3)
        for figure_path in figure_paths:
            png = Path(figure_path).with_suffix(".png")
            image_path = png if png.exists() else Path(figure_path).with_suffix(".tiff")
            if not image_path.exists():
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(image_path), width=Inches(6.55))
            stem = Path(figure_path).stem
            add_caption(doc, FIGURE_CAPTIONS.get(stem, f"Fig. {len(doc.inline_shapes)}. {stem.replace('_', ' ')}."))

    def reference_profile(path: str | Path | None) -> dict[str, Any]:
        profile: dict[str, Any] = {
            "page_width": Cm(21.0),
            "page_height": Cm(29.7),
            "top_margin": Inches(0.787),
            "bottom_margin": Inches(0.787),
            "left_margin": Inches(0.787),
            "right_margin": Inches(0.787),
            "header_distance": Inches(0.591),
            "footer_distance": Inches(0.0),
            "body_font": "Times New Roman",
            "body_size": 12.0,
            "title_size": 16.0,
        }
        if not path:
            return profile
        ref_path = Path(path)
        if not ref_path.exists():
            return profile
        try:
            ref_doc = Document(str(ref_path))
            ref_sec = ref_doc.sections[0]
            for key in [
                "page_width",
                "page_height",
                "top_margin",
                "bottom_margin",
                "left_margin",
                "right_margin",
                "header_distance",
                "footer_distance",
            ]:
                profile[key] = getattr(ref_sec, key)
            normal_style = ref_doc.styles["Normal"]
            if normal_style.font.name:
                profile["body_font"] = normal_style.font.name
            if normal_style.font.size:
                profile["body_size"] = float(normal_style.font.size.pt)
            try:
                title_style = ref_doc.styles["Title"]
                if title_style.font.size:
                    profile["title_size"] = float(title_style.font.size.pt)
            except Exception:
                pass
        except Exception:
            return profile
        return profile

    tokens = reference_profile(format_reference_docx)
    _ = format_reference_docx
    doc = Document()
    if doc.paragraphs:
        for paragraph in list(doc.paragraphs):
            p = paragraph._element
            p.getparent().remove(p)
    for table in list(doc.tables):
        tbl = table._element
        tbl.getparent().remove(tbl)

    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = tokens["page_width"]
    section.page_height = tokens["page_height"]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, tokens[attr])
    section.header_distance = tokens["header_distance"]
    section.footer_distance = tokens["footer_distance"]

    normal = doc.styles["Normal"]
    normal.font.name = tokens["body_font"]
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), tokens["body_font"])
    normal.font.size = Pt(tokens["body_size"])
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.15

    lines = manuscript.splitlines()
    i = 0
    table_idx = 1
    inserted_figures = False
    in_references = False
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("## VIII.") and not inserted_figures:
            add_figures(doc)
            inserted_figures = True
        if stripped.startswith("# "):
            add_heading(doc, stripped[2:].strip(), 1)
        elif stripped.startswith("## "):
            heading_text = stripped[3:].strip()
            in_references = heading_text.lower().startswith("references")
            add_heading(doc, heading_text, 2)
        elif stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), 3)
        elif stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(parse_table_line(lines[i]))
                i += 1
            add_markdown_table(doc, table_lines, table_idx)
            table_idx += 1
            continue
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style=None)
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.first_line_indent = Cm(-0.2)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(stripped)
            style_run(run, size=12.0)
        else:
            text = stripped.replace("`", "")
            if in_references:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                run = p.add_run(text)
                style_run(run, size=10.0)
            elif text.startswith("Fig.") or text.startswith("Table "):
                add_caption(doc, text)
            else:
                add_body_paragraph(doc, text)
        i += 1
    if not inserted_figures:
        add_figures(doc)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    return output


def build_dataset_profile(rows: list[dict[str, Any]], *, human_checked_rows: list[dict[str, Any]] | None = None) -> dict[str, Any]:
    human_checked_rows = human_checked_rows or []
    evidence_counts = [len(row.get("evidence_spans") or []) for row in rows]
    reasoning_counts = [len(row.get("reasoning_spans") or []) for row in rows]
    fact_counts = [len(row.get("facts") or []) for row in rows]
    claim_counts = [len(row.get("claims") or []) for row in rows]
    label_counts = Counter(str(row.get("liability_label") or "unknown") for row in rows)
    project_counts = Counter(str(row.get("project_type") or "unknown") for row in rows)
    dispute_counts = Counter(str(row.get("dispute_type") or "unknown") for row in rows)
    return {
        "num_cases": len(rows),
        "human_checked_cases": len(human_checked_rows),
        "liability_label_counts": dict(label_counts),
        "project_type_counts": dict(project_counts.most_common(12)),
        "dispute_type_counts": dict(dispute_counts.most_common(12)),
        "evidence_count_mean": round(mean(evidence_counts), 3) if evidence_counts else 0.0,
        "evidence_count_median": round(median(evidence_counts), 3) if evidence_counts else 0.0,
        "evidence_count_std": round(pstdev(evidence_counts), 3) if len(evidence_counts) > 1 else 0.0,
        "evidence_count_min": min(evidence_counts) if evidence_counts else 0,
        "evidence_count_max": max(evidence_counts) if evidence_counts else 0,
        "reasoning_count_mean": round(mean(reasoning_counts), 3) if reasoning_counts else 0.0,
        "fact_count_mean": round(mean(fact_counts), 3) if fact_counts else 0.0,
        "claim_count_mean": round(mean(claim_counts), 3) if claim_counts else 0.0,
        "evidence_counts": evidence_counts,
    }


def build_complete_manuscript(
    *,
    profile: dict[str, Any],
    metrics: dict[str, Any],
    calibration_report: dict[str, Any],
    calibration_metrics: dict[str, Any],
    token_usage: dict[str, Any],
    runtime: dict[str, Any],
    references: list[str],
    figure_paths: list[Path],
) -> str:
    result_block = formal_results_block(metrics)
    calibration_block = calibration_results_block(calibration_report, calibration_metrics)
    data_block = data_profile_table(profile)
    label_block = label_distribution_table(profile)
    figure_block = figure_reference_block(figure_paths)
    metric_table = metric_definition_table()
    baseline_table = baseline_design_table()
    runtime_block = runtime_cost_block(token_usage, runtime)
    if metrics:
        results_context = (
            "The completed Stage 3 run is treated as the formal 500-case pilot evidence for this draft. "
            "Because candidate labels are machine-assisted and only a 50-case subset has been human checked, "
            "the classification metrics are interpreted as weak-supervised pilot indicators rather than as performance "
            "against a fully validated human reference set. Baseline and ablation runs remain part of the planned "
            "controlled comparison and should be reported when those experiment files are generated."
        )
        figure5_sentence = (
            "Fig. 5 reports the formal 500-case results dashboard generated from the current metrics file."
        )
        conclusion_sentence = (
            "The current harness now contains a complete 500-case full MAD-RA pilot run, refreshed result tables, "
            "and an updateable figure package; subsequent reruns can replace the metrics and regenerate the same manuscript artifacts."
        )
    else:
        results_context = (
            "The manuscript is intentionally structured so that formal model results can be updated by rerunning the pilot "
            "workflow and regenerating figures. Until the formal 500-case predictions exist, all model-comparison claims "
            "remain placeholders and are not interpreted as empirical findings.\n\n"
            "The current draft therefore separates three empirical layers. The first layer is real data descriptives from "
            "the 500-case pilot dataset. The second layer is calibration evidence from the 30-case Qwen run, which evaluates "
            "output validity and prompt robustness but not final performance. The third layer is the pending formal "
            "model-comparison evidence, which will become the Results section once Stage 3 is rerun. This separation prevents "
            "calibration findings from being overstated as substantive model performance.\n\n"
            "When the formal run is completed, the Results narrative should focus on four comparisons. First, full MAD-RA "
            "should be compared with single-agent and RAG-only baselines on classification and evidence metrics. Second, "
            "the no-evidence-verification ablation should be used to isolate unsupported-claim control. Third, the "
            "no-coordination and single-round ablations should be used to isolate convergence and re-deliberation. Fourth, "
            "the controlled single-agent baselines should be used to address the token-budget and context-length alternative explanation."
        )
        figure5_sentence = "Fig. 5 is the results dashboard that will update when formal 500-case metrics are available."
        conclusion_sentence = (
            "Once the formal 500-case pilot is rerun, the results tables and figures can be refreshed without changing "
            "the paper's conceptual structure."
        )

    manuscript = f"""# MAD-RA: A Blackboard-Mediated and Evidence-Grounded Multi-Agent Deliberation Framework for Responsibility Attribution in Construction Delay Disputes

## Abstract
Responsibility attribution in construction delay disputes is not an ordinary delay-prediction or text-classification task. It requires deliberation over factual statements, party claims, contractual duties, causal delay chains, evidence sufficiency, and competing responsibility narratives. Existing construction document analytics, RAG systems, and LLM-based decision-support tools improve information extraction and response generation, but they rarely provide an explicit mechanism for multi-party argumentation, evidence verification, conflict detection, re-deliberation, and convergence control. This paper proposes MAD-RA, a blackboard-mediated and evidence-grounded multi-agent deliberation framework for responsibility attribution in construction delay disputes. MAD-RA represents each dispute as a CaseRecord with stable evidence identifiers, maintains a Shared Case Blackboard containing the case profile, evidence registry, claim ledger, argument messages, critique messages, verification records, conflict graph, consensus state, revision history, and audit log, and coordinates six role agents through structured ClaimMessage, EvidenceMessage, ChallengeMessage, RebuttalMessage, ConcessionMessage, VerificationMessage, and CoordinatorFeedback objects. The framework computes label disagreement, allocation disagreement, evidence disagreement, unsupported-claim rate, composite disagreement, consensus score, and stop conditions, and uses targeted re-deliberation when conflicts arise in labels, allocations, evidence references, causality explanations, or notice-duty interpretations. The study design combines 10-case mock debugging, 30-case API calibration, a {profile['num_cases']}-case weak-supervised pilot using machine-assisted candidate labels, a {profile['human_checked_cases']}-case human-checked subset, controlled baselines, ablations, stability analysis, and cost/runtime diagnostics. MAD-RA outputs responsibility labels, allocation vectors, evidence chains, unresolved conflict points, unsupported claims, evidence-preparation suggestions, negotiation focus, and preventive management implications. The framework is a decision-support system for project managers, contract engineers, and claim engineers; it is not legal advice, not an automated adjudication system, and not a substitute for adjudicators.

## Index Terms
Construction delay disputes; responsibility attribution; multi-agent systems; large language models; evidence grounding; contract management; engineering management; decision support.

## Managerial Relevance Statement
Construction managers, contract administrators, project controls teams, and claim engineers often face responsibility questions before a dispute reaches formal adjudication. MAD-RA supports these users by transforming unstructured dispute records into evidence-indexed deliberation artifacts. The framework highlights which responsibility claims are supported by evidence, which claims remain unsupported, where owner and contractor interpretations diverge, and which records should be prepared for negotiation or expert review. Its value is managerial rather than adjudicative: it improves evidence readiness, auditability, and negotiation discipline while preserving human accountability.

## 1. Introduction
Schedule delay is one of the most persistent sources of conflict in construction projects. Once delay escalates into a claim or formal dispute, the managerial question changes from whether the project was delayed to who should bear responsibility, which events are excusable or compensable, and what evidence supports a proposed allocation [1]-[17]. This shift matters because responsibility attribution is not reducible to a single critical-path calculation. It depends on how contract duties, notices, approvals, instructions, recovery measures, payment events, site records, and concurrent delay arguments interact within a contested project history.

Traditional delay-analysis and claims-management research provides the technical and legal vocabulary for this task. Studies on probabilistic delay analysis, concurrent delay, claim presentation, and change-dispute outcome prediction clarify how delay responsibility can be reasoned about in principle [1]-[17], [34]-[52]. Yet these methods often rely on expert interpretation, project-specific scheduling evidence, or manually curated cases. They are less suited to large collections of adjudication texts where the goal is to reuse adjudication logic as structured decision support.

Artificial intelligence has created new opportunities for construction text analytics. NLP and machine-learning models have been used for contract risk and responsibility assessment, design-build requirement categorization, contract summarization, construction dispute precedent summarization, RAG-based construction information retrieval, and claim-report generation [18]-[33]. These studies show that engineering management knowledge embedded in documents can be extracted, retrieved, and reorganized. However, most systems still produce a retrieved answer, a classification label, or a generated report rather than a deliberative responsibility judgement with explicit disagreement and evidence-grounding controls.

Responsibility attribution in delay disputes has three features that make ordinary single-agent generation insufficient. First, the task is multi-perspective: owner-side and contractor-side narratives often interpret the same event differently. Second, the task is evidence-sensitive: a plausible claim is not useful if it cannot be traced to contemporaneous evidence. Third, the task is convergence-sensitive: managers need to know whether a conclusion is stable across role interpretations or merely the product of a single prompt. These features motivate a framework that can expose conflict, check support, and report uncertainty.

MAD-RA addresses this need by treating responsibility attribution as evidence-grounded multi-agent deliberation. The framework does not ask a language model to produce one long answer. Instead, it structures the task into role-specific outputs with a strict schema, validates evidence identifiers, computes disagreement over liability labels, allocation estimates, and evidence references, and triggers re-deliberation when disagreement remains high. The final output is a decision-support package containing the responsibility conclusion, allocation estimate, evidence IDs, unsupported-claim rate, consensus score, conflict points, and management implications.

This paper contributes to engineering management research in four ways. First, at the task level, it reformulates construction delay responsibility attribution as an evidence-grounded multi-agent deliberation problem rather than a conventional label-prediction or single-agent question-answering problem. Second, at the method level, it introduces a blackboard-mediated communication protocol in which agents read and write structured messages against a shared case state rather than freely chatting. Third, at the algorithm level, it defines computable disagreement, unsupported-claim rate, consensus score, stop condition, and targeted re-deliberation mechanisms. Fourth, at the management level, it frames responsibility attribution as decision support for evidence preparation, negotiation focus, claim-file readiness, and dispute prevention rather than as automated adjudication.

The managerial motivation is equally important. Delay disputes are not isolated legal events; they are symptoms of weak information governance, inconsistent claim preparation, and misaligned responsibility narratives. A system that only predicts a label gives managers limited help because it does not show which records are missing, which assertions are unsupported, or which disagreement points should shape negotiation. MAD-RA therefore treats responsibility attribution as an organizational learning problem. Historical adjudication texts are used to learn how responsibility arguments are structured, while the resulting outputs are designed to improve future evidence collection, claim-file readiness, and preventive project controls.

The paper is organized as follows. Section 2 reviews delay responsibility, construction document analytics, multi-agent LLMs, and communication mechanisms. Section 3 formulates the technical optimization problem. Section 4 presents the blackboard-mediated MAD-RA methodology. Section 5 describes the experimental design. Section 6 reports pilot findings. Sections 7 and 8 discuss implications and limitations. Section 9 concludes.

## 2. Related Work
### 2.1 Construction Delay Responsibility and Claim Management
Delay analysis research has established a rich set of concepts for understanding responsibility in construction disputes. Integrated probabilistic delay analysis, standardized delay-claim management, design responsibility allocation, change-dispute outcome prediction, scheduling practice, global claims, payment disputes, change-order negotiation, and concurrent-delay analysis all inform how responsibility is assessed [1]-[17]. Classic delay and disruption literature further emphasizes claim preparation, presentation, extension of time, loss and expense, and project-specific causation [34]-[52].

This literature is essential for MAD-RA because it supplies the domain logic that role agents should consider: causation, critical-path effect, notice compliance, concurrent delay, mitigation, compensability, and evidence sufficiency. Its limitation is not conceptual weakness but scalability and process representation. Most studies do not convert large adjudication corpora into structured evidence spans, nor do they model the argumentation process through which opposing responsibility narratives are reconciled.

### 2.2 NLP/RAG for Construction Documents and Claims
NLP-based construction contract analysis has moved from information extraction toward risk and responsibility assessment [18]-[21]. This line of work is particularly relevant because it shows that risk-bearing parties, contractual obligations, and responsibility cues can be computationally identified. Explainability is central in this context: responsibility-sensitive outputs must be inspectable by managers and experts.

However, contract review and dispute responsibility attribution differ in their input structure and decision logic. Contract review usually examines clauses before performance failure. Delay-dispute attribution examines after-the-fact narratives that mix facts, claims, evidence, adjudicative reasoning, and final decisions. A method that identifies risky clauses cannot directly determine how responsibility should be allocated when evidence is incomplete or conflicting.

Recent construction AI studies show that retrieval-augmented generation can improve access to construction management knowledge, safety information, claim documents, and dispute precedents [22]-[33]. RAG is valuable for MAD-RA because responsibility reasoning must be tied to stable evidence units. Evidence-grounded reasoning studies in NLP also show why generated conclusions should be connected to explicit evidence chains [60], [61].

RAG alone remains insufficient for responsibility attribution. A retrieved passage may support one party's claim, weaken another party's claim, or merely provide background. Responsibility attribution requires comparing these roles, not only retrieving text. MAD-RA therefore treats retrieval as the evidence-support layer and adds role-specific reasoning, evidence verification, disagreement scoring, and coordination.

### 2.3 LLM-Based Multi-Agent Systems
Multi-agent LLM systems have been applied to manufacturing scheduling, construction meeting issue mitigation, intelligent geotechnical design, healthcare record modelling, legal judgment prediction, and evidence-grounded legal reasoning [25]-[27], [53]-[62]. These studies motivate the use of specialized agents and deliberation for complex decision tasks. They also show that multi-stage reasoning can improve transparency when outputs are structured and audited.

The research gap is domain-specific and process-specific. Existing systems rarely focus on construction delay responsibility attribution, and even fewer combine role-based dispute reasoning with stable evidence identifiers, unsupported-claim controls, and convergence metrics. MAD-RA is designed to fill this gap by aligning the multi-agent architecture with the structure of real delay disputes: owner view, contractor view, delay-causality view, evidence sufficiency, and coordination.

### 2.4 Communication Mechanisms in Multi-Agent Decision Support
Multi-agent systems differ not only by agent roles but also by communication protocols. A set of agents that independently generates answers and concatenates them is not equivalent to a deliberative system. In engineering decision support, communication must maintain a single source of case facts, track which evidence supports which claim, expose local conflicts, control re-deliberation, and preserve an audit trail. The design inspiration from construction delay-response strategy generation is therefore translated into MAD-RA as four formal mechanisms: a Shared Case Blackboard, structured message passing, coordinator-mediated deliberation, and BPMN-style workflow traceability. Unlike a generic memory buffer, the blackboard is an evidence-indexed case state with typed ledgers and graphs. Unlike informal agent chat, message passing uses typed objects with evidence IDs and required actions. Unlike a direct central judge, the coordinator computes conflict and convergence diagnostics and controls the next deliberation task. Unlike a visual BPMN drawing, the workflow trace is serialized as audit data for reproducibility and review.

Table-based prediction studies and LLM-based document systems also differ in their treatment of uncertainty. In many construction-management applications, uncertainty is reported as model confidence or statistical error. In responsibility attribution, uncertainty also has an evidentiary meaning. A conclusion may be uncertain because the model is unstable, because the candidate label is noisy, because evidence is missing, or because both parties contributed to delay. MAD-RA separates these risks by reporting disagreement score, consensus score, unsupported-claim rate, and evidence-chain overlap. This separation is necessary for managerial use because each risk implies a different action: collect missing records, invite expert review, reframe the negotiation issue, or downgrade the conclusion to insufficient evidence.

The literature therefore supports four design principles. First, the framework must preserve domain concepts from delay-analysis and claims research. Second, it must enforce traceability from claims to evidence spans. Third, it must expose disagreement rather than hide it in a single generated narrative. Fourth, it must make communication and workflow state auditable. These principles guide the problem formulation and methodology.

## 3. Technical Problem Formulation
### 3.1 Construction Delay Responsibility Attribution as Evidence-Grounded Deliberation
Construction delay responsibility attribution is formulated here as an evidence-grounded deliberation problem rather than a one-step prediction problem. A case contains factual statements, party claims, evidence spans, adjudicative reasoning spans, and a candidate decision outcome. The computational task is not simply to assign a label to the entire text. It is to construct a responsibility attribution that is supported by valid evidence IDs, consistent with role-specific interpretations, and useful for managerial review.

Each case is represented as a CaseRecord containing case ID, project type, dispute type, parties, facts, claims, evidence spans, reasoning spans, final decision, candidate liability label, and weak-supervised allocation reference. Evidence spans receive stable identifiers so that evidence precision, recall, and Hit@k can be evaluated. The CaseRecord is defined as:

\\[
x=(F,C,E,R,D)
\\]

where \\(F\\) denotes factual statements, \\(C\\) denotes party claims, \\(E\\) denotes evidence spans, \\(R\\) denotes adjudicative reasoning spans, and \\(D\\) denotes the decision outcome or candidate label. The role-agent set is \\(A=\\{{a_1,\\ldots,a_m\\}}\\). The objective is to produce a decision-support output \\(O=(y,p,E^*,K,U,S,H)\\), where \\(y\\) is a liability label, \\(p\\) is a normalized allocation vector, \\(E^*\\) is the evidence chain, \\(K\\) is the unresolved conflict set, \\(U\\) is unsupported-claim rate, \\(S\\) is consensus score, and \\(H\\) is the audit trail.

The candidate label set contains owner responsibility, contractor responsibility, shared responsibility, insufficient evidence, and other. Allocation is represented as a party-level percentage vector when available. Because allocation labels in historical adjudication documents can be noisy or implicit, allocation is treated conservatively: exact numeric equality is not the only target, and allocation variance is reported as a stability metric rather than as a stand-alone legal conclusion.

### 3.2 Technical Limitations of One-Shot and Free-Chat Multi-Agent Reasoning
One-shot LLM reasoning is weak for this task because it collapses owner-side, contractor-side, causal, contractual, and evidentiary interpretations into a single generated answer. Retrieval-augmented generation improves access to evidence, but retrieval alone does not decide whether a cited span supports, contradicts, weakly supports, or fails to support a responsibility claim. Free-chat multi-agent systems expose multiple voices, but they often lack a shared state, typed messages, evidence-indexed claim ledgers, and auditable stop conditions.

The technical limitation is therefore not only insufficient language understanding. It is insufficient coordination. If every agent answers independently, conflicts may remain hidden. If every agent is forced to answer in every round, token cost rises even when only one conflict type needs attention. If stopping is based on a fixed number of rounds or informal coordinator judgement, convergence is not measurable. A technical contribution must therefore specify how the coordinator chooses what to revise, which agents to activate, which evidence to verify, and when to stop.

### 3.3 Evidence-Constrained Deliberation Optimization Problem
MAD-RA treats deliberation as a constrained trajectory optimization problem. At round \\(t\\), the coordinator observes the shared blackboard state, the conflict graph, verification records, and the current consensus state. The goal is to search for a responsibility output that improves consensus, evidence grounding, contractual consistency, and managerial usefulness while reducing disagreement, unsupported claims, and runtime/token cost.

The overall quality objective is:

\\[
\\max O^* =
\\alpha S
+\\beta G
+\\gamma R
-\\delta D
-\\eta U
-\\mu C
\\]

where \\(S\\) is consensus score, \\(G\\) is evidence-grounding quality, \\(R\\) is contract/rule consistency, \\(D\\) is composite disagreement, \\(U\\) is unsupported-claim rate, and \\(C\\) is token/runtime cost. The coefficients are non-negative pilot parameters and should be examined through sensitivity analysis rather than treated as universal constants. This objective makes the technical target explicit: MAD-RA optimizes deliberation quality under evidence and cost constraints, not merely classification accuracy.

### 3.4 Adaptive Coordination Policy
The coordinator action at round \\(t\\) is defined as:

\\[
a_t=(A_t,K_t,E_t^+,\\Theta_t)
\\]

where \\(A_t\\subseteq A\\) is the subset of agents triggered in the next deliberation step, \\(K_t\\) is the conflict-point set to be addressed, \\(E_t^+\\) is the evidence set requiring supplementation or verification, and \\(\\Theta_t\\) contains the current thresholds and weighting parameters. This definition turns coordination into a computable policy rather than a narrative instruction.

The adaptive re-deliberation policy is:

\\[
A_t^*=\\arg\\max_{{A_t\\subseteq A}}\\Delta Q(A_t,K_t)-\\lambda Cost(A_t)
\\]

where \\(\\Delta Q(A_t,K_t)\\) is the expected improvement in deliberation quality when a subset of agents addresses the current conflict points, and \\(Cost(A_t)\\) is the expected token/runtime cost. In the current prototype, this policy is implemented as a rule-based coordinator over the conflict graph. Label conflict triggers owner and contractor perspectives; allocation conflict triggers delay-causality and contract-rule reasoning; evidence conflict triggers evidence verification; notice-duty conflict triggers the contract-rule agent; and high unsupported-claim rate prevents unsupported responsibility attribution from being promoted to the final explanation.

This adaptive policy is the main technical distinction between MAD-RA and ordinary multi-agent prompting. The framework does not assume that more agents and more rounds are always better. It uses conflict-specific diagnostics to decide whether further deliberation is justified and which part of the responsibility argument should be revised.

### 3.5 Research Questions and Hypotheses
This paper addresses four technical research questions.

RQ1: Does the Shared Case Blackboard improve responsibility attribution and evidence grounding compared with multi-agent reasoning without a shared state?

RQ2: Does targeted re-deliberation improve convergence and cost-quality tradeoff compared with fixed-round all-agent re-deliberation?

RQ3: Does adaptive coordination reduce unsupported-claim rate and evidence disagreement compared with no-coordination and no-evidence-verification ablations?

RQ4: Does the additional token/runtime cost of blackboard-mediated deliberation translate into evidence grounding, stability, auditability, and managerial usefulness?

The technical problem is therefore not merely to generate a responsibility label, but to optimize an evidence-constrained deliberation trajectory. MAD-RA addresses this problem by introducing an adaptive coordinator that operates over a shared blackboard, identifies conflict-specific deficiencies, selectively triggers role agents, updates the conflict graph, and stops deliberation only when disagreement, unsupported claims, and consensus satisfy explicit thresholds.

## 4. Methodology
### 4.1 CaseRecord Construction
MAD-RA starts from a structured CaseRecord. The purpose of this representation is to avoid treating a dispute file as an undifferentiated block of text. Facts, party claims, evidence spans, reasoning spans, and final decisions are separated so that each part has a clear role in inference and evaluation. Evidence spans are the only valid source of evidence IDs in model outputs. Reasoning spans may provide context but are not accepted as evidence IDs unless explicitly converted into evidence spans.

### 4.2 Shared Case Blackboard
MAD-RA uses a Shared Case Blackboard rather than independent agent memories. The blackboard is the single evidence-indexed case state:

\\[
B_t=\\{{X,E_t,G_t,M_t,V_t,C_t,H_t\\}}
\\]

where \\(X\\) is the CaseRecord, \\(E_t\\) is the evidence registry, \\(G_t\\) is the claim-evidence-responsibility graph, \\(M_t\\) is the message pool, \\(V_t\\) is the verification record set, \\(C_t\\) is the conflict graph, and \\(H_t\\) is audit history. The blackboard also stores claim ledger, consensus state, and revision history. This mechanism prevents each agent from reconstructing a separate version of the dispute and enables workflow-level auditability.

### 4.3 Structured Message Passing
Agents communicate through typed messages rather than free-form chat. The supported message types are ClaimMessage, EvidenceMessage, ChallengeMessage, RebuttalMessage, ConcessionMessage, VerificationMessage, and CoordinatorFeedback. Each message includes message_id, round_id, sender, receiver, claim_id, evidence_ids, message_type, content, confidence, and required_action. Role-agent message generation is defined as:

\\[
m_i^{{(t)}}=f_i(B_t,\\pi_i)
\\]

where \\(i\\) indexes the agent and \\(\\pi_i\\) is the role-specific policy. The term policy is used deliberately: the framework defines analytical roles and message obligations, not personalities.

### 4.4 Role-Specific Deliberation Agents
MAD-RA uses six roles. The Owner Perspective Agent analyzes owner instructions, design changes, approval delay, payment delay, site handover, and suspension instruction. The Contractor Perspective Agent analyzes resource allocation, construction organization, subcontractor coordination, progress control, notice compliance, and mitigation effort. The Delay Causality Agent analyzes delay events, critical-path effects, concurrent delay, causal chains, and mitigation effects. The Contract Rule Agent analyzes extension of time, liquidated damages, variations, notice, force majeure, compensability, and contractual risk allocation. The Evidence Verification Agent checks claim-evidence links, missing evidence, weak support, contradictory evidence, and unsupported claims. The Coordinator Agent detects conflicts, posts targeted feedback, controls convergence, synthesizes final responsibility attribution, and exports the audit trail.

All agents output the same JSON schema: agent role, liability label, allocation, key claims, evidence IDs, reasoning steps, uncertainty, and unsupported claims. This uniform schema makes outputs comparable. It also prevents the framework from relying on unstructured prose when computing disagreement.

### 4.5 Evidence Verification and Claim-Evidence Graph
The evidence-verification mechanism is a first-class component. It checks whether key claims are supported by valid evidence IDs and identifies missing, weak, or conflicting evidence. The deterministic final guard removes unsupported final attribution claims when no valid evidence ID supports them. If no valid evidence remains for a responsibility conclusion, the framework must move toward an insufficient-evidence output rather than a fluent but unsupported attribution.

The verification function is:

\\[
v(c_j,e_k)\\in\\{{support,contradict,weak,missing\\}}
\\]

where \\(c_j\\) is a claim and \\(e_k\\) is an evidence span. The argument graph uses claim, evidence, responsibility, delay-event, and party nodes with supports, contradicts, weakly_supports, missing, causes, mitigates, and allocates_to edges.

### 4.6 Conflict Detection and Targeted Re-Deliberation
The coordinator computes disagreement across responsibility labels, allocation vectors, evidence references, and unsupported claims. Label disagreement is:

\\[
D_y^{{(t)}}=1-\\frac{{\\max_y |\\{{i:y_i^{{(t)}}=y\\}}|}}{{|A|}}
\\]

Allocation disagreement is:

\\[
D_p^{{(t)}}=\\frac{{2}}{{|A|(|A|-1)}}\\sum_{{i<j}}\\frac{{\\|p_i^{{(t)}}-p_j^{{(t)}}\\|_1}}{{2}}
\\]

where \\(p_i\\) is a normalized responsibility allocation vector whose elements sum to 1. Evidence disagreement is:

\\[
D_e^{{(t)}}=1-\\frac{{2}}{{|A|(|A|-1)}}\\sum_{{i<j}}J(E_i^{{(t)}},E_j^{{(t)}})
\\]

where \\(J\\) is Jaccard similarity and \\(J(E_i,E_j)=1\\) when \\(|E_i\\cup E_j|=0\\). Unsupported-claim rate is:

\\[
U^{{(t)}}=\\frac{{|unsupported\\ claims|}}{{\\max(1,|key\\ claims|)}}
\\]

Composite disagreement is:

\\[
D^{{(t)}}=w_yD_y^{{(t)}}+w_pD_p^{{(t)}}+w_eD_e^{{(t)}}+w_uU^{{(t)}},\\quad \\sum w=1
\\]

The weights are not theoretical constants. The pilot setting uses a default configuration and the experimental design includes sensitivity analysis to test whether conclusions depend on the weighting scheme.

If disagreement exceeds the threshold, the coordinator generates conflict points and asks relevant agents to re-deliberate around those points. The process stops when disagreement is below threshold and consensus is above threshold, or when the maximum number of rounds is reached. This design turns convergence into a measurable property rather than an implicit judgement by the model.

### 4.7 Consensus, Stop Condition, and Audit Trail
The stop condition is:

\\[
Stop\\ if\\ \\left(D^{{(t)}}\\le\\tau_D\\right)\\land\\left(U^{{(t)}}\\le\\tau_U\\right)\\land\\left(S^{{(t)}}\\ge\\tau_S\\right)\\quad or\\quad t=T.
\\]

Each round is recorded as a workflow trace: task_started -> agent_argument_submitted -> evidence_verified -> conflict_detected -> feedback_posted -> re_deliberation_completed -> final_output_generated. The implementation exports this as audit_log.jsonl and embeds the workflow trace in each prediction JSON object. This is workflow-level auditability, not merely a process diagram.

### 4.8 Management Decision-Support Output
The final output includes managerial implications: key risk points, evidence-preparation suggestions, claim-negotiation focus, and preventive actions. This design is important for IEEE TEM positioning. The framework does not claim to decide legal responsibility. It helps managers understand evidence readiness, negotiation priorities, and dispute-prevention actions.

### F. Complexity and Cost Considerations
For a case with A role agents, R deliberation rounds, and K retrieved evidence spans, the dominant API cost scales with A times R plus any coordination calls. The current prototype uses standard-library keyword retrieval and deterministic coordination by default, which keeps the pipeline reproducible and auditable. Multi-agent deliberation increases token cost relative to a single-agent baseline, so the pilot reports token usage, runtime, retry rate, malformed-output rate, and schema compliance.

The framework deliberately separates two kinds of coordination. The first is deterministic coordination, which computes disagreement and evidence overlap from structured outputs. The second is optional LLM coordination, which can write a richer synthesis but must still pass the same schema and evidence guard. The deterministic layer is essential for reproducibility because it ensures that key metrics do not depend entirely on generative phrasing. This design also makes ablation interpretable: removing the evidence-verification agent, coordination, or re-deliberation changes a specific mechanism rather than the entire pipeline.

The stop condition is defined by two thresholds. If disagreement is lower than tau and consensus is higher than gamma, the deliberation stops. Otherwise, the coordinator generates conflict-focused feedback. In practice, this feedback names the disputed liability labels, divergent allocation estimates, or non-overlapping evidence references. The next round must respond to those points rather than restart from a blank prompt. This is the mechanism that distinguishes MAD-RA from ordinary parallel prompting.

### Algorithm 1: Blackboard-mediated MAD-RA Deliberation
Input: CaseRecord \\(x\\), agent set \\(A\\), maximum rounds \\(T\\), thresholds \\(\\tau_D,\\tau_U,\\tau_S\\).

Output: responsibility label, allocation vector, evidence chain, conflict points, consensus score, audit log, and management implications.

1. Initialize blackboard \\(B_0\\) from \\(x\\).
2. Retrieve and register evidence spans in \\(E_0\\).
3. For each round \\(t=1,\\ldots,T\\):
4. Coordinator posts task to the blackboard.
5. Role agents read \\(B_t\\) and submit ClaimMessages.
6. EvidenceVerificationAgent verifies claim-evidence links.
7. Coordinator computes disagreement and unsupported-claim rate.
8. Coordinator updates the conflict graph.
9. If the stop condition is satisfied, break.
10. Otherwise, generate targeted feedback and trigger local re-deliberation.
11. Generate the final decision-support output.
12. Export the audit trail.

## 5. Experimental Design
The local dataset contains parsed construction-related adjudication documents. The formal pilot subset contains {profile['num_cases']} CaseRecord items. Labels are machine-assisted candidate labels or weak-supervised labels. They are not treated as fully human-validated labels. A {profile['human_checked_cases']}-case human-checked subset is used for small-scale validation of liability labels, allocations, evidence spans, unsupported claims, and explanation quality.

{data_block}

{label_block}

The experiment is organized into seven layers. First, a 10-case mock debugging stage checks schema compliance, structured message passing, audit-log serialization, and evidence-ID validity. Second, a 30-case Qwen API calibration stage checks JSON validity, retry rate, unsupported-claim rate, malformed-output rate, evidence grounding, and prompt robustness. Third, a 500-case pilot uses machine-assisted candidate labels and weak labels; it is not a fully validated benchmark. Fourth, a 50-case human-checked subset evaluates liability plausibility, allocation plausibility, evidence support, explanation usefulness, and unsupported-claim detection. Fifth, baselines and ablations compare single-agent same-context, single-agent long-prompt, RAG-only, multi-agent without blackboard, multi-agent direct-message-only, no evidence verification, no coordinator, single-round MAD-RA, all-agent fixed-round deliberation, and full MAD-RA. The no-blackboard ablation is critical because it tests whether the blackboard communication mechanism contributes beyond simply calling multiple agents. The all-agent fixed-round baseline is critical because it tests whether adaptive coordination and targeted re-deliberation add value beyond repeated all-agent prompting. Sixth, stability analysis samples 100 representative cases from the pilot and runs each case three times. Seventh, cost/runtime analysis reports tokens, runtime, estimated API cost, retry rate, and malformed JSON rate.

{calibration_block}

{baseline_table}

{metric_table}

Evaluation metrics are grouped into four families. Responsibility attribution performance uses Accuracy, Macro-F1, and per-class F1. Evidence grounding quality uses Evidence Precision, Evidence Recall, Hit@k, unsupported-claim rate, and evidence-chain overlap. Deliberation behavior uses number of rounds, disagreement reduction, consensus score, conflict resolution rate, and stop-condition satisfaction rate. Practical usability uses runtime, token cost, human usefulness rating, evidence-preparation usefulness, and negotiation-focus usefulness.

The human-checked subset is not used to claim full validation of all candidate labels. Its purpose is narrower: it checks whether a small stratified sample has plausible liability labels, allocation direction, evidence spans, unsupported-claim flags, and explanation quality. This distinction is important for research integrity. The 500-case pilot is a weak-supervised empirical study, while the 50-case subset provides a small human validation anchor.

The controlled single-agent and fixed-round baselines are included to answer likely reviewer concerns. A multi-agent system may appear better simply because it uses more tokens, sees more context, or receives a longer instruction. The single-agent same-context baseline holds the retrieved evidence input close to MAD-RA. The single-agent long-prompt baseline exposes a compact role checklist in one prompt. The all-agent fixed-round baseline keeps all role agents active for a fixed number of rounds, which isolates the value of adaptive conflict-specific coordination. If full MAD-RA improves evidence grounding or stability beyond these baselines, the result is more plausibly attributable to the deliberation architecture.

## 6. Results
{result_block['main_text']}

{runtime_block}

{results_context}

## 7. Discussion
The visual argument is organized around six figures. Fig. 1 explains the blackboard-mediated architecture. Fig. 2 formalizes structured message passing and shared case state. Fig. 3 presents conflict detection and targeted re-deliberation. Fig. 4 reports the pilot protocol and validation layers. {figure5_sentence} Fig. 6 translates audit records and final outputs into management decision-support actions.

{figure_block}

MAD-RA makes three conceptual shifts. First, it shifts from delay prediction to responsibility attribution. Second, it shifts from unstructured explanation to evidence-indexed argumentation. Third, it shifts from one-shot model output to deliberation with measurable disagreement and convergence. These shifts matter for engineering management because responsibility decisions affect evidence preservation, negotiation posture, claim strategy, and dispute escalation.

The pilot design also clarifies what should and should not be claimed. If MAD-RA improves Macro-F1 but produces unsupported claims, it is not managerially reliable. If it improves evidence precision and stability with only modest label gains, the framework may still be valuable because managers need traceable and auditable reasoning. Conversely, if multi-agent deliberation raises token cost without improving evidence grounding or stability, the added complexity should be questioned. This is why the experiment reports evidence metrics, stability metrics, and cost metrics alongside classification performance.

The role of the evidence-verification agent is especially important. In construction disputes, unsupported but plausible narratives can be harmful because they may distort claim preparation or negotiation. By forcing key claims to cite evidence IDs and reporting unsupported-claim rate, MAD-RA makes evidence gaps visible rather than hiding them behind confident prose.

The management output extends the framework beyond post-hoc adjudication analysis. The same evidence-grounded logic can help project teams maintain delay-event registers, identify missing notices, prepare claim files, structure negotiation agendas, and prevent disputes from escalating. This is the strongest IEEE TEM contribution: the method connects AI reasoning to engineering management workflows.

A second managerial implication concerns accountability. Managers should not use MAD-RA as a substitute for professional judgement. Instead, they can use it as a structured checklist for asking better questions: Which party narrative is supported? Which evidence IDs are decisive? Which claims are unsupported? Where do role agents disagree? Which additional records would reduce uncertainty? These questions are practical, auditable, and aligned with the way claim files are prepared.

Finally, MAD-RA can support organizational learning. Repeated analysis of historical cases can reveal recurring weak points in project controls, such as late design approvals, incomplete notices, undocumented recovery plans, or poor linkage between schedule updates and correspondence. These insights can be fed back into contract administration protocols and evidence-governance routines.

## 8. Limitations
The current dataset uses machine-assisted candidate labels and weak-supervised labels. The {profile['human_checked_cases']}-case subset supports small-scale human validation, but it does not establish full validation for all {profile['num_cases']} cases. The formal Results section must therefore avoid language implying a complete human-validated reference set.

The current retrieval layer uses standard-library keyword retrieval. This choice makes the prototype easy to reproduce, but it may miss semantically relevant evidence. Future versions can introduce vector retrieval or knowledge-graph retrieval without changing the CaseRecord or AgentOutput schema.

The framework is currently calibrated on adjudication texts rather than live project records. Transferring it to ongoing projects requires additional work on document ingestion, privacy, contractual jurisdiction, and human review. The system should remain a decision-support framework, not a legal advice system or adjudication substitute.

Finally, multi-agent reasoning increases runtime and token cost. The tradeoff is justified only when it improves evidence grounding, stability, interpretability, or human usefulness. Cost and latency are therefore not implementation details; they are part of the managerial evaluation.

Ethical and legal boundaries also require explicit treatment. The system processes legal and contractual text, but its output remains analytical support. It should not be used to advise parties on legal rights, determine liability, or replace arbitrators, judges, lawyers, or delay experts. The appropriate use case is internal analysis, claim preparation, negotiation planning, and expert-review support. Any deployment should include data privacy controls, jurisdiction-specific review, and clear user-facing disclaimers.

## 9. Conclusion
This paper presents MAD-RA, an evidence-grounded multi-agent deliberation framework for responsibility attribution in construction delay disputes. The framework uses structured CaseRecords, role-specific agents, evidence verification, disagreement scoring, re-deliberation, and management decision-support output to make responsibility reasoning more auditable. The current draft provides a complete IEEE TEM-oriented manuscript structure, data profile, figure package, and results-update harness. {conclusion_sentence}

## Appendix A. Reproducibility Checklist
- Dataset format: JSONL with one CaseRecord per line.
- Label terminology: machine-assisted candidate labels and weak-supervised labels.
- Human validation: small-scale human-checked subset only.
- Main model: full MAD-RA.
- Baselines: no evidence verification, no coordination, single round, single agent, single-agent same context, single-agent long prompt, and RAG-only.
- Stability test: 100 representative cases, three repeated runs.
- Required metrics: Accuracy, Macro-F1, per-class F1, evidence precision, evidence recall, Hit@k, unsupported-claim rate, label consistency, allocation variance, evidence-chain overlap, consensus score mean and standard deviation.
- Cost reporting: tokens per case, runtime per case, API calls per case, retry rate, malformed-output rate, and schema compliance.

## References
{chr(10).join(references)}
"""
    return manuscript


def formal_results_block(metrics: dict[str, Any]) -> dict[str, str]:
    if not metrics:
        return {
            "abstract_sentence": f"{FORMAL_RESULTS_MARKER} The formal 500-case model-comparison results are pending and will be inserted after the Qwen pilot is rerun.",
            "main_text": (
                f"{FORMAL_RESULTS_MARKER}\n\n"
                "The formal 500-case predictions have not yet been generated in this workspace. "
                "Accordingly, this Results section reports the real data profile and calibration checks, "
                "but it does not claim model superiority. After rerunning Stage 3, replace this block with "
                "the generated accuracy, Macro-F1, evidence metrics, stability metrics, ablation results, "
                "and cost diagnostics."
            ),
        }
    items = [
        ("num_predictions", "N"),
        ("accuracy", "Accuracy"),
        ("macro_f1", "Macro-F1"),
        ("evidence_precision", "Evidence precision"),
        ("evidence_recall", "Evidence recall"),
        ("hit@5", "Hit@5"),
        ("unsupported_claim_rate", "Unsupported-claim rate"),
        ("label_consistency", "Label consistency"),
        ("allocation_variance", "Allocation variance"),
        ("evidence_chain_overlap", "Evidence-chain overlap"),
        ("consensus_score_mean", "Mean consensus score"),
    ]
    rows = []
    for key, label in items:
        if key in metrics:
            rows.append(f"| {label} | {format_number(metrics[key])} |")
    table = "| Metric | Formal 500-case value |\n|---|---|\n" + "\n".join(rows)
    return {
        "abstract_sentence": "The formal 500-case pilot reports the main model metrics described in the Results section.",
        "main_text": "The formal 500-case pilot produced the following metrics.\n\n" + table,
    }


def calibration_results_block(report: dict[str, Any], metrics: dict[str, Any]) -> str:
    merged = {**metrics, **report}
    if not merged:
        return "The 30-case Qwen calibration report is not available in the current output directory."
    keys = [
        ("successful_json_rate", "Successful JSON rate"),
        ("schema_compliance_rate", "Schema compliance rate"),
        ("retry_rate", "Retry rate"),
        ("malformed_output_rate", "Malformed-output rate"),
        ("unsupported_claim_rate_mean", "Unsupported-claim rate mean"),
        ("evidence_precision", "Evidence precision"),
        ("evidence_recall", "Evidence recall"),
        ("hit@5", "Hit@5"),
    ]
    rows = [f"| {label} | {format_number(merged[key])} |" for key, label in keys if key in merged]
    return "Table III reports the Stage 2 Qwen calibration checks.\n\n| Calibration metric | Value |\n|---|---|\n" + "\n".join(rows)


def data_profile_table(profile: dict[str, Any]) -> str:
    rows = [
        ("Cases", profile["num_cases"]),
        ("Human-checked subset", profile["human_checked_cases"]),
        ("Mean evidence spans per case", profile["evidence_count_mean"]),
        ("Median evidence spans per case", profile["evidence_count_median"]),
        ("Evidence span range", f"{profile['evidence_count_min']} to {profile['evidence_count_max']}"),
        ("Mean reasoning spans per case", profile["reasoning_count_mean"]),
        ("Mean facts per case", profile["fact_count_mean"]),
        ("Mean claims per case", profile["claim_count_mean"]),
    ]
    body = "\n".join(f"| {name} | {value} |" for name, value in rows)
    return "Table I summarizes the pilot dataset.\n\n| Dataset property | Value |\n|---|---|\n" + body


def label_distribution_table(profile: dict[str, Any]) -> str:
    total = max(1, int(profile["num_cases"]))
    rows = []
    for label, count in profile["liability_label_counts"].items():
        rows.append(f"| {label} | {count} | {count / total:.3f} |")
    return "Table II reports the candidate liability-label distribution.\n\n| Candidate label | Count | Share |\n|---|---:|---:|\n" + "\n".join(rows)


def baseline_design_table() -> str:
    return """Table IV defines the baseline and ablation design.

| Mode | Purpose |
|---|---|
| Full MAD-RA | Complete evidence-grounded multi-agent deliberation |
| No evidence verification | Tests marginal value of the evidence-verification agent |
| No coordination | Tests marginal value of disagreement detection and convergence |
| Single round | Tests whether re-deliberation adds value |
| Single agent | Tests ordinary one-agent responsibility attribution |
| Single-agent same context | Controls for the same retrieved evidence input |
| Single-agent long prompt | Controls for prompt length and role checklist exposure |
| RAG-only | Tests retrieval without role-based deliberation |
| Multi-agent without blackboard | Tests whether multiple agents alone are sufficient |
| Multi-agent direct-message only | Tests point-to-point message passing without a shared case state |
| All-agent fixed rounds | Tests repeated all-agent prompting without adaptive targeted re-deliberation |
| No blackboard | Key communication ablation for the Shared Case Blackboard |
"""


def metric_definition_table() -> str:
    return """Table V defines the evaluation metrics.

| Metric family | Metrics | Interpretation |
|---|---|---|
| Responsibility attribution | Accuracy, Macro-F1, per-class F1 | Responsibility-label performance against candidate labels |
| Evidence grounding | Evidence precision, evidence recall, Hit@k, unsupported-claim rate, evidence-chain overlap | Whether cited evidence IDs support claims |
| Deliberation behavior | Number of rounds, disagreement reduction, consensus score, conflict resolution rate, stop-condition satisfaction rate | Whether deliberation converges |
| Practical usability | Runtime, token cost, human usefulness rating, evidence-preparation usefulness, negotiation-focus usefulness | Feasibility and management value |
"""


def runtime_cost_block(token_usage: dict[str, Any], runtime: dict[str, Any]) -> str:
    if not token_usage and not runtime:
        return "Runtime, token, and cost diagnostics will be updated after the formal 500-case pilot is rerun."
    rows = []
    for key, label in [
        ("average_tokens_per_case", "Average tokens per case"),
        ("average_api_calls_per_case", "Average API calls per case"),
        ("retry_rate", "Retry rate"),
        ("malformed_output_rate", "Malformed-output rate"),
    ]:
        if key in token_usage:
            rows.append(f"| {label} | {format_number(token_usage[key])} |")
    if "average_running_time_per_case" in runtime:
        rows.append(f"| Average runtime per case | {format_number(runtime['average_running_time_per_case'])} seconds |")
    return "| Runtime or cost metric | Value |\n|---|---:|\n" + "\n".join(rows)


def figure_reference_block(paths: list[Path]) -> str:
    if not paths:
        return "No figures have been generated."
    rows = []
    for idx, path in enumerate(paths, 1):
        rows.append(f"| Fig. {idx} | `{path.name}` | Editable SVG/PDF exported by Python harness |")
    return "| Figure | File | Note |\n|---|---|---|\n" + "\n".join(rows)


def make_tem_figures(
    *,
    output_dir: str | Path,
    profile: dict[str, Any],
    metrics: dict[str, Any],
    calibration: dict[str, Any],
    token_usage: dict[str, Any],
    runtime: dict[str, Any],
) -> list[Path]:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyArrowPatch, Rectangle

    plt.rcParams.update(
        {
            "font.family": "serif",
            "font.serif": ["Times New Roman", "Times", "DejaVu Serif"],
            "svg.fonttype": "none",
            "pdf.fonttype": 42,
            "font.size": 8,
            "axes.linewidth": 0.6,
            "axes.spines.right": False,
            "axes.spines.top": False,
        }
    )
    output = Path(output_dir)
    output.mkdir(parents=True, exist_ok=True)
    paths: list[Path] = []

    def save(fig, name: str) -> Path:
        base = output / name
        fig.savefig(base.with_suffix(".svg"), bbox_inches="tight")
        fig.savefig(base.with_suffix(".pdf"), bbox_inches="tight")
        fig.savefig(base.with_suffix(".png"), dpi=300, bbox_inches="tight")
        try:
            fig.savefig(base.with_suffix(".tiff"), dpi=600, bbox_inches="tight")
        except Exception:
            pass
        plt.close(fig)
        return base.with_suffix(".svg")

    def add_box(ax, xy, text, w=1.7, h=0.55, fc="#F7F7F7"):
        rect = Rectangle(xy, w, h, facecolor=fc, edgecolor="#666666", linewidth=0.7)
        ax.add_patch(rect)
        ax.text(xy[0] + w / 2, xy[1] + h / 2, text, ha="center", va="center", fontsize=8, wrap=True)
        return rect

    def add_arrow(ax, start, end):
        ax.add_patch(FancyArrowPatch(start, end, arrowstyle="->", mutation_scale=8, linewidth=0.7, color="#333333"))

    fig, ax = plt.subplots(figsize=(7.2, 4.0))
    ax.axis("off")
    add_box(ax, (0.1, 2.7), "CaseRecord\nfacts claims evidence")
    for y, label in [(3.4, "Owner\nagent"), (2.6, "Contractor\nagent"), (1.8, "Delay analysis\nagent"), (1.0, "Evidence verification\nagent")]:
        add_box(ax, (2.2, y), label, w=1.55, h=0.45)
        add_arrow(ax, (1.8, 2.98), (2.2, y + 0.22))
        add_arrow(ax, (3.75, y + 0.22), (4.5, 2.35))
    add_box(ax, (4.5, 2.0), "Coordination\nscore disagreement", w=1.7, h=0.7, fc="#EEF3F8")
    add_box(ax, (6.8, 2.0), "Decision-support\noutput", w=1.7, h=0.7, fc="#F5F1E8")
    add_arrow(ax, (6.2, 2.35), (6.8, 2.35))
    ax.text(0.1, 3.95, "Fig. 1. MAD-RA architecture", fontsize=10, fontweight="bold")
    ax.set_xlim(0, 8.8)
    ax.set_ylim(0.6, 4.2)
    paths.append(save(fig, "fig1_madra_architecture"))

    fig, ax = plt.subplots(figsize=(7.2, 4.0))
    ax.axis("off")
    steps = [
        ("Preprocess\nCaseRecord", 0.2),
        ("Retrieve\nEvidence IDs", 1.7),
        ("Role\nargumentation", 3.2),
        ("Compute\ndisagreement", 4.7),
        ("Stop or\nre-deliberate", 6.2),
    ]
    for label, x in steps:
        add_box(ax, (x, 2.1), label, w=1.15, h=0.65)
    for (_, x1), (_, x2) in zip(steps, steps[1:]):
        add_arrow(ax, (x1 + 1.15, 2.43), (x2, 2.43))
    add_arrow(ax, (6.75, 2.1), (3.75, 1.25))
    add_arrow(ax, (3.75, 1.25), (3.75, 2.1))
    ax.text(0.2, 3.45, "Fig. 2. Deliberation, disagreement scoring, and stop condition", fontsize=10, fontweight="bold")
    ax.text(4.85, 1.15, "if disagreement > tau", fontsize=8, color="#666666")
    ax.set_xlim(0, 7.8)
    ax.set_ylim(0.8, 3.8)
    paths.append(save(fig, "fig2_deliberation_workflow"))

    fig, ax = plt.subplots(figsize=(7.2, 3.7))
    ax.axis("off")
    stages = [
        ("Stage 1", "Mock 10\nschema and JSONL"),
        ("Stage 2", "Qwen 30\ncalibration"),
        ("Stage 3", "Formal 500\npilot results"),
        ("Stability", "100 cases\n3 runs"),
    ]
    for i, (title, subtitle) in enumerate(stages):
        x = 0.35 + i * 1.75
        add_box(ax, (x, 1.7), f"{title}\n{subtitle}", w=1.35, h=0.8, fc=["#F7F7F7", "#EEF3F8", "#F5F1E8", "#F8EEEE"][i])
        if i < len(stages) - 1:
            add_arrow(ax, (x + 1.35, 2.1), (x + 1.75, 2.1))
    ax.text(0.35, 3.05, "Fig. 3. Three-stage SCI pilot protocol", fontsize=10, fontweight="bold")
    ax.set_xlim(0, 7.5)
    ax.set_ylim(1.0, 3.3)
    paths.append(save(fig, "fig3_pilot_protocol"))

    fig, axes = plt.subplots(1, 2, figsize=(7.2, 3.6))
    labels = list(profile["liability_label_counts"].keys())
    counts = list(profile["liability_label_counts"].values())
    axes[0].barh(range(len(labels)), counts, color="#4C78A8", edgecolor="#333333", linewidth=0.4)
    axes[0].set_yticks(range(len(labels)), labels)
    axes[0].set_xlabel("Cases")
    axes[0].set_title("Candidate label distribution", fontsize=9, fontweight="bold")
    evidence_counts = profile.get("evidence_counts", [])
    axes[1].hist(evidence_counts, bins=min(12, max(3, len(set(evidence_counts)))), color="#59A14F", edgecolor="#333333", linewidth=0.4)
    axes[1].set_xlabel("Evidence spans per case")
    axes[1].set_ylabel("Cases")
    axes[1].set_title("Evidence-span distribution", fontsize=9, fontweight="bold")
    fig.suptitle("Fig. 4. Formal 500-case pilot data profile", x=0.02, y=1.02, ha="left", fontsize=10, fontweight="bold")
    paths.append(save(fig, "fig4_dataset_profile"))

    fig, ax = plt.subplots(figsize=(7.2, 3.8))
    ax.axis("off")
    dashboard = [
        ("Accuracy", metrics.get("accuracy")),
        ("Macro-F1", metrics.get("macro_f1")),
        ("Evidence P", metrics.get("evidence_precision") or calibration.get("evidence_precision")),
        ("Evidence R", metrics.get("evidence_recall") or calibration.get("evidence_recall")),
        ("Hit@5", metrics.get("hit@5") or calibration.get("hit@5")),
        ("Unsupported rate", metrics.get("unsupported_claim_rate") or calibration.get("unsupported_claim_rate_mean")),
    ]
    for i, (name, value) in enumerate(dashboard):
        x = 0.35 + (i % 3) * 2.2
        y = 2.55 - (i // 3) * 1.1
        add_box(ax, (x, y), f"{name}\n{format_number(value) if value is not None else 'pending'}", w=1.75, h=0.75, fc="#F7F7F7")
    ax.text(0.35, 3.55, "Fig. 5. Evaluation dashboard for formal 500-case pilot", fontsize=10, fontweight="bold")
    note = (
        "Metrics generated from the Stage 3 full MAD-RA 500-case pilot."
        if metrics
        else "Pending cells are updated after Stage 3 formal 500-case predictions are generated."
    )
    ax.text(0.35, 0.85, note, fontsize=8, color="#666666")
    ax.set_xlim(0, 7.2)
    ax.set_ylim(0.6, 3.8)
    paths.append(save(fig, "fig5_results_dashboard"))

    fig, ax = plt.subplots(figsize=(7.2, 3.8))
    ax.axis("off")
    loop = [
        ("Risk points", 0.3, 2.4),
        ("Evidence\npreparation", 2.2, 2.4),
        ("Negotiation\nfocus", 4.1, 2.4),
        ("Preventive\nactions", 3.0, 1.2),
        ("Project learning", 1.1, 1.2),
    ]
    for label, x, y in loop:
        add_box(ax, (x, y), label, w=1.45, h=0.65, fc="#F5F1E8")
    arrows = [((1.75, 2.72), (2.2, 2.72)), ((3.65, 2.72), (4.1, 2.72)), ((4.7, 2.4), (4.0, 1.85)), ((3.0, 1.52), (2.55, 1.52)), ((1.8, 1.85), (0.95, 2.4))]
    for start, end in arrows:
        add_arrow(ax, start, end)
    ax.text(0.3, 3.35, "Fig. 6. Management decision-support loop", fontsize=10, fontweight="bold")
    ax.set_xlim(0, 6.2)
    ax.set_ylim(0.8, 3.6)
    paths.append(save(fig, "fig6_management_loop"))

    source_data = {
        "profile": {k: v for k, v in profile.items() if k != "evidence_counts"},
        "metrics": metrics,
        "calibration": calibration,
        "token_usage": token_usage,
        "runtime": runtime,
        "figures": [str(path) for path in paths],
    }
    (output / "figure_source_data.json").write_text(json.dumps(source_data, ensure_ascii=False, indent=2), encoding="utf-8")
    return paths


def make_tem_figures(
    *,
    output_dir: str | Path,
    profile: dict[str, Any],
    metrics: dict[str, Any],
    calibration: dict[str, Any],
    token_usage: dict[str, Any],
    runtime: dict[str, Any],
) -> list[Path]:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyArrowPatch, Rectangle

    plt.rcParams.update(
        {
            "font.family": "serif",
            "font.serif": ["Times New Roman", "Times", "DejaVu Serif"],
            "svg.fonttype": "none",
            "pdf.fonttype": 42,
            "font.size": 8,
            "axes.linewidth": 0.6,
            "axes.spines.right": False,
            "axes.spines.top": False,
        }
    )
    output = Path(output_dir)
    output.mkdir(parents=True, exist_ok=True)
    paths: list[Path] = []

    def save(fig, name: str) -> Path:
        base = output / name
        fig.savefig(base.with_suffix(".svg"), bbox_inches="tight")
        fig.savefig(base.with_suffix(".pdf"), bbox_inches="tight")
        fig.savefig(base.with_suffix(".png"), dpi=300, bbox_inches="tight")
        try:
            fig.savefig(base.with_suffix(".tiff"), dpi=600, bbox_inches="tight")
        except Exception:
            pass
        plt.close(fig)
        return base.with_suffix(".svg")

    def box(ax, x, y, text, w=1.35, h=0.45, fc="#F7F7F7", ec="#555555", fs=7.5):
        rect = Rectangle((x, y), w, h, facecolor=fc, edgecolor=ec, linewidth=0.7)
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fs, wrap=True)
        return rect

    def arrow(ax, start, end, color="#333333"):
        ax.add_patch(FancyArrowPatch(start, end, arrowstyle="->", mutation_scale=8, linewidth=0.7, color=color))

    # Fig. 1: architecture
    fig, ax = plt.subplots(figsize=(7.2, 4.1))
    ax.axis("off")
    box(ax, 0.15, 2.25, "CaseRecord\nfacts claims evidence", w=1.45, h=0.65)
    box(ax, 2.15, 1.3, "Shared Case Blackboard\ncase profile | evidence registry\nclaim ledger | message pool\nconflict graph | audit log", w=2.25, h=1.65, fc="#EEF3F8", fs=7)
    roles = [
        ("Owner\nPerspective", 0.35, 3.35),
        ("Contractor\nPerspective", 2.0, 3.35),
        ("Delay\nCausality", 3.65, 3.35),
        ("Contract\nRule", 5.3, 3.35),
        ("Evidence\nVerification", 5.3, 0.55),
        ("Coordinator", 3.65, 0.55),
    ]
    for label, x, y in roles:
        box(ax, x, y, label, w=1.15, h=0.55, fc="#F7F7F7")
        arrow(ax, (x + 0.57, y), (3.25, 2.95 if y > 2 else 1.3))
        arrow(ax, (3.25, 2.95 if y > 2 else 1.3), (x + 0.57, y + 0.55), color="#6A6A6A")
    arrow(ax, (1.6, 2.58), (2.15, 2.58))
    box(ax, 5.65, 1.65, "Decision-support\nJSON output", w=1.25, h=0.75, fc="#F5F7FA")
    arrow(ax, (4.4, 2.1), (5.65, 2.02))
    ax.text(0.15, 4.05, "Fig. 1. Blackboard-mediated MAD-RA architecture", fontsize=10, fontweight="bold")
    ax.set_xlim(0, 7.1)
    ax.set_ylim(0.35, 4.3)
    paths.append(save(fig, "fig1_madra_architecture"))

    # Fig. 2: structured messages and blackboard
    fig, ax = plt.subplots(figsize=(7.2, 3.9))
    ax.axis("off")
    message_types = ["Claim", "Evidence", "Challenge", "Rebuttal", "Concession", "Verification", "Feedback"]
    for i, label in enumerate(message_types):
        x = 0.2 + (i % 4) * 1.55
        y = 2.45 - (i // 4) * 0.85
        box(ax, x, y, f"{label}\nMessage", w=1.2, h=0.55, fc="#F7F7F7")
        arrow(ax, (x + 0.6, y), (3.55, 1.55))
    box(ax, 2.35, 0.65, "message_id | round_id | sender | receiver\nclaim_id | evidence_ids | type | content\nconfidence | required_action", w=3.1, h=0.75, fc="#EEF3F8", fs=7)
    box(ax, 5.75, 1.15, "Blackboard\nsingle source\nof case state", w=1.15, h=1.05, fc="#F5F7FA", fs=7)
    arrow(ax, (5.45, 1.05), (5.75, 1.55))
    ax.text(0.2, 3.55, "Fig. 2. Structured message passing over shared case state", fontsize=10, fontweight="bold")
    ax.set_xlim(0, 7.1)
    ax.set_ylim(0.35, 3.8)
    paths.append(save(fig, "fig2_deliberation_workflow"))

    # Fig. 3: conflict and re-deliberation
    fig, ax = plt.subplots(figsize=(7.2, 3.8))
    ax.axis("off")
    steps = [
        ("Aggregate\nmessages", 0.2, 2.2),
        ("Detect\nconflicts", 1.65, 2.2),
        ("Update\nconflict graph", 3.1, 2.2),
        ("Compute\nD, U, S", 4.55, 2.2),
        ("Stop\ncondition", 6.0, 2.2),
    ]
    for label, x, y in steps:
        box(ax, x, y, label, w=1.05, h=0.6, fc="#F7F7F7")
    for (_, x1, y1), (_, x2, y2) in zip(steps, steps[1:]):
        arrow(ax, (x1 + 1.05, y1 + 0.3), (x2, y2 + 0.3))
    box(ax, 1.2, 0.95, "label | allocation | evidence\ncausality | notice duty", w=1.8, h=0.7, fc="#EEF3F8", fs=7)
    box(ax, 3.55, 0.95, "targeted feedback\nonly conflicted claims", w=1.7, h=0.7, fc="#F5F7FA", fs=7)
    arrow(ax, (6.55, 2.2), (4.4, 1.65))
    arrow(ax, (3.55, 1.3), (1.2, 1.3))
    ax.text(0.2, 3.35, "Fig. 3. Conflict graph and targeted re-deliberation", fontsize=10, fontweight="bold")
    ax.text(5.65, 1.7, "if thresholds not met", fontsize=7, color="#666666")
    ax.set_xlim(0, 7.2)
    ax.set_ylim(0.65, 3.6)
    paths.append(save(fig, "fig3_pilot_protocol"))

    # Fig. 4: experiment protocol and baselines
    fig, ax = plt.subplots(figsize=(7.2, 4.0))
    ax.axis("off")
    stages = [
        ("Mock 10", "schema\nmessages\naudit log"),
        ("Qwen 30", "JSON\nretry\ngrounding"),
        ("Pilot 500", "weak labels\nfull MAD-RA\nmetrics"),
        ("Human 50", "plausibility\nusefulness\nsupport"),
    ]
    for i, (title, subtitle) in enumerate(stages):
        x = 0.25 + i * 1.72
        box(ax, x, 2.55, f"{title}\n{subtitle}", w=1.25, h=0.9, fc="#F7F7F7")
        if i < len(stages) - 1:
            arrow(ax, (x + 1.25, 3.0), (x + 1.72, 3.0))
    baselines = ["same context", "long prompt", "RAG-only", "no blackboard", "direct message", "no verifier", "no coordinator", "single round"]
    for i, label in enumerate(baselines):
        x = 0.35 + (i % 4) * 1.65
        y = 1.35 - (i // 4) * 0.55
        box(ax, x, y, label, w=1.25, h=0.35, fc="#EEF3F8", fs=6.8)
    ax.text(0.25, 3.75, "Fig. 4. Pilot study protocol and controlled communication ablations", fontsize=10, fontweight="bold")
    ax.set_xlim(0, 7.2)
    ax.set_ylim(0.45, 4.0)
    paths.append(save(fig, "fig4_dataset_profile"))

    # Fig. 5: result dashboard
    fig, ax = plt.subplots(figsize=(7.2, 3.8))
    ax.axis("off")
    dashboard = [
        ("Accuracy", metrics.get("accuracy")),
        ("Macro-F1", metrics.get("macro_f1")),
        ("Evidence P", metrics.get("evidence_precision") or calibration.get("evidence_precision")),
        ("Evidence R", metrics.get("evidence_recall") or calibration.get("evidence_recall")),
        ("Hit@5", metrics.get("hit@5") or calibration.get("hit@5")),
        ("Unsupported rate", metrics.get("unsupported_claim_rate") or calibration.get("unsupported_claim_rate_mean")),
    ]
    for i, (name, value) in enumerate(dashboard):
        x = 0.35 + (i % 3) * 2.2
        y = 2.55 - (i // 3) * 1.1
        box(ax, x, y, f"{name}\n{format_number(value) if value is not None else 'pending'}", w=1.75, h=0.75, fc="#F7F7F7", fs=8)
    ax.text(0.35, 3.55, "Fig. 5. Evaluation dashboard for formal 500-case pilot", fontsize=10, fontweight="bold")
    note = "Metrics generated from the Stage 3 full MAD-RA 500-case pilot." if metrics else "Pending cells are updated after Stage 3 predictions."
    ax.text(0.35, 0.85, note, fontsize=8, color="#666666")
    ax.set_xlim(0, 7.2)
    ax.set_ylim(0.6, 3.8)
    paths.append(save(fig, "fig5_results_dashboard"))

    # Fig. 6: audit and management output
    fig, ax = plt.subplots(figsize=(7.2, 4.0))
    ax.axis("off")
    workflow = ["task", "argument", "verify", "conflict", "feedback", "final"]
    for i, label in enumerate(workflow):
        x = 0.25 + i * 1.1
        box(ax, x, 2.75, label, w=0.82, h=0.4, fc="#F7F7F7", fs=7)
        if i < len(workflow) - 1:
            arrow(ax, (x + 0.82, 2.95), (x + 1.1, 2.95))
    outputs = [
        ("evidence chain", 0.45, 1.55),
        ("unsupported claims", 2.05, 1.55),
        ("conflict points", 3.65, 1.55),
        ("evidence prep", 0.45, 0.85),
        ("negotiation focus", 2.05, 0.85),
        ("preventive actions", 3.65, 0.85),
    ]
    for label, x, y in outputs:
        box(ax, x, y, label, w=1.25, h=0.45, fc="#EEF3F8", fs=7)
    box(ax, 5.45, 1.05, "decision-support\nnot adjudication", w=1.35, h=0.85, fc="#F5F7FA", fs=7)
    arrow(ax, (4.9, 1.1), (5.45, 1.45))
    ax.text(0.25, 3.55, "Fig. 6. Workflow-level auditability and management outputs", fontsize=10, fontweight="bold")
    ax.set_xlim(0, 7.1)
    ax.set_ylim(0.55, 3.85)
    paths.append(save(fig, "fig6_management_loop"))

    source_data = {
        "profile": {k: v for k, v in profile.items() if k != "evidence_counts"},
        "metrics": metrics,
        "calibration": calibration,
        "token_usage": token_usage,
        "runtime": runtime,
        "figures": [str(path) for path in paths],
    }
    (output / "figure_source_data.json").write_text(json.dumps(source_data, ensure_ascii=False, indent=2), encoding="utf-8")
    return paths


def build_figure_contract(*, metrics: dict[str, Any]) -> str:
    status = "formal results available" if metrics else "formal results pending"
    return f"""# Figure Contract

Backend: Python matplotlib only.

Style: AIC/AEI/SCI engineering informatics style; white background; Times New Roman or serif fallback; thin borders; black/grey/muted blue palette; editable SVG/PDF exports.

Core conclusion: MAD-RA is a blackboard-mediated decision-support framework that improves auditability by combining shared case state, structured messages, evidence verification, disagreement scoring, targeted re-deliberation, and workflow traceability.

Evidence chain:
- Fig. 1: blackboard-mediated architecture claim.
- Fig. 2: structured message protocol claim.
- Fig. 3: conflict graph and targeted re-deliberation claim.
- Fig. 4: pilot, baseline, and validation design claim.
- Fig. 5: formal result dashboard, status = {status}.
- Fig. 6: auditability and management decision-support claim.

Review risks:
- Do not imply fully human-validated labels.
- Do not present pending metrics as empirical results.
- Do not frame MAD-RA as legal advice or adjudication replacement.
"""


def build_ai_figure_prompts() -> str:
    return """# AI Drawing Prompts for Schematic Redrawing

Use these prompts only for schematic refinement. Keep all text editable if redrawn in Illustrator, PowerPoint, Figma, or vector AI tools.

## Prompt 1: MAD-RA Architecture
Create a clean engineering informatics schematic on a white background. Show a CaseRecord input block on the left with facts, claims, evidence spans, and reasoning spans. In the center, show a Shared Case Blackboard containing case profile, evidence registry, claim ledger, message pool, conflict graph, consensus state, revision history, and audit log. Around it, show six agents: owner perspective, contractor perspective, delay causality, contract rule, evidence verification, and coordinator. Use thin grey lines, muted blue accents, Times New Roman-style labels, no gradients, no icons, no 3D effects.

## Prompt 2: Deliberation Loop
Create a flow diagram showing structured message passing over a shared case blackboard. Include ClaimMessage, EvidenceMessage, ChallengeMessage, RebuttalMessage, ConcessionMessage, VerificationMessage, and CoordinatorFeedback. Each message should show compact fields: message_id, round_id, sender, receiver, claim_id, evidence_ids, type, content, confidence, required_action. Style as an engineering informatics paper figure with white background, thin black borders, muted blue highlights, and compact typography.

## Prompt 3: Managerial Decision-Support Loop
Create a workflow-level auditability figure showing task_started, agent_argument_submitted, evidence_verified, conflict_detected, feedback_posted, re_deliberation_completed, and final_output_generated. Connect the workflow to final management outputs: key evidence chain, unsupported claims, unresolved conflict points, evidence preparation suggestions, negotiation focus, and preventive management implications. Show that MAD-RA supports managers and does not replace adjudication. Use restrained black, grey, and muted blue colors.
"""


def build_ppt_absorption_checklist() -> str:
    return """# PPT Mechanism Absorption Checklist

Source concept: multi-agent delay-response strategy generation slides 19-21.

Absorbed into MAD-RA as follows:

| PPT mechanism | MAD-RA research mechanism | Implementation artifact | Paper location |
|---|---|---|---|
| Shared blackboard for progress, cost, risk, and draft plans | Shared Case Blackboard as single evidence-indexed case state | `madra/blackboard.py` (`BlackboardState`, `EvidenceRegistry`, `ClaimLedger`, `ConflictGraph`, `ConsensusState`, `AuditLog`) | Methodology 4.2 |
| Point-to-point request/concession/explanation messages | Structured Message Passing with typed messages and evidence IDs | `madra/messages.py` (`ClaimMessage`, `EvidenceMessage`, `ChallengeMessage`, `RebuttalMessage`, `ConcessionMessage`, `VerificationMessage`, `CoordinatorFeedback`) | Methodology 4.3 |
| Central coordination agent for multi-round meetings | Coordinator-Mediated Deliberation for conflict detection, targeted feedback, convergence control, and final synthesis | `madra/protocol.py`, `madra/coordinator.py`, `madra/pipeline.py` | Methodology 4.6-4.7 |
| BPMN workflow recording triggers and outputs | BPMN-style workflow trace and JSONL audit log for workflow-level auditability | `madra/workflow.py`, `AuditLog.export_jsonl()` | Methodology 4.7 |
| LLM upper layer + CPM/EVM/rules lower layer | Role-specific deliberation plus rule-aware contract agent and evidence graph | `contract_rule_agent`, `ArgumentGraph`, `EvidenceVerificationAgent` | Methodology 4.4-4.5 |

The PPT diagrams were not copied. They were translated into formal SCI paper mechanisms and redrawn as black/grey/muted-blue vector figures.
"""


def build_update_manifest() -> str:
    return """# Rerun And Update Manifest

Run these commands after API access is available again.

```powershell
python run_pilot_study.py --output-dir outputs/pilot_500 --limit 500 --model qwen-plus --temperature 0 --max-rounds 2
python run_tem_harness.py --dataset outputs/pilot_500/dataset_madra_500.jsonl --output-dir outputs/ieee_tem_harness
```

If you want only to refresh figures after metrics files already exist:

```powershell
python run_tem_harness.py --dataset outputs/pilot_500/dataset_madra_500.jsonl --output-dir outputs/ieee_tem_harness --no-docx
```

Update targets:
- `metrics_500.json`
- `token_usage_500.json`
- `runtime_500.json`
- `cost_500.json`
- `stage3_qwen_500_predictions.jsonl`
- `outputs/ieee_tem_harness/figures/*`
- `outputs/ieee_tem_harness/IEEE_TEM_MADRA_full_draft.md`
"""


def read_json(path: str | Path | None) -> dict[str, Any]:
    if not path:
        return {}
    p = Path(path)
    if not p.exists():
        return {}
    return json.loads(p.read_text(encoding="utf-8"))


def format_number(value: Any) -> str:
    if value is None:
        return "pending"
    if isinstance(value, int):
        return str(value)
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    if abs(number) >= 100:
        return f"{number:.1f}"
    return f"{number:.3f}"

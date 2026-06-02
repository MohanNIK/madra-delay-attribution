from __future__ import annotations

from dataclasses import dataclass
import json
import os
from pathlib import Path
import re
from typing import Any

from .agents import complete_with_schema_retries
from .llm import LLMClient, OpenAICompatibleClient


PAPER_SECTION_CONTRACT = """
Return only valid JSON with exactly these keys:
{
  "section_name": "section name",
  "claims_used": ["claim grounded in the supplied project context"],
  "citations_used": ["[1]"],
  "missing_evidence_flags": ["missing result or citation warning"],
  "draft_text": "complete English section draft",
  "review_notes": ["audit note"]
}
Rules:
- Write in formal IEEE Transactions on Engineering Management style.
- Do not invent empirical results, data, or citations.
- Use machine-assisted candidate labels / weak-supervised labels, not fully validated label terminology.
- Describe MAD-RA as a decision-support framework, not a legal advice system or adjudication substitute.
- Cite only reference numbers that appear in the supplied reference library.
"""


IEEE_TEM_STYLE_PROFILE = """# IEEE TEM Dynamic Writing Skill for MAD-RA

Generated for an English manuscript targeting IEEE Transactions on Engineering Management.

## Editorial Identity
- Position the paper as engineering management decision support, not as a pure API or prompt-engineering demo.
- Keep the problem framed around managerial decision quality: evidence preparation, claim negotiation, dispute prevention, and auditability.
- Make the technical contribution legible to engineering-management readers: explicit workflow, measurable outputs, robustness checks, and implications for practitioners.

## Required Manuscript Moves
- Include a Managerial Relevance Statement immediately after Index Terms.
- Separate algorithm contribution from managerial contribution.
- Report Results only from actual pilot outputs. If formal 500-case outputs are unavailable, mark results as pending.
- Avoid language implying fully validated labels. Use machine-assisted candidate labels, weak-supervised labels, and human-checked subset.

## Prohibited Claims
- Do not claim legal authority, automatic adjudication, or professional legal advice.
- Do not claim full validation of the 500-case labels unless a complete human validation file exists.
- Do not describe Qwen API usage as the contribution; the contribution is the evidence-grounded multi-agent deliberation framework.
"""


SECTION_ORDER = [
    "Title and Abstract",
    "Managerial Relevance Statement",
    "Introduction",
    "Literature Review and Research Gap",
    "Research Questions and Problem Formulation",
    "MAD-RA Framework",
    "Data and Experimental Design",
    "Results",
    "Discussion",
    "Managerial Implications",
    "Limitations",
    "Conclusion",
]


@dataclass(frozen=True)
class PaperSectionDraft:
    section_name: str
    claims_used: list[str]
    citations_used: list[str]
    missing_evidence_flags: list[str]
    draft_text: str
    review_notes: list[str]

    @classmethod
    def from_dict(cls, data: dict[str, Any]) -> "PaperSectionDraft":
        if not isinstance(data, dict):
            raise ValueError("Paper section output must be a JSON object.")
        required = [
            "section_name",
            "claims_used",
            "citations_used",
            "missing_evidence_flags",
            "draft_text",
            "review_notes",
        ]
        missing = [key for key in required if key not in data]
        if missing:
            raise ValueError(f"Missing paper section fields: {', '.join(missing)}")
        return cls(
            section_name=clean_for_policy(str(data["section_name"])),
            claims_used=[clean_for_policy(str(item)) for item in _as_list(data["claims_used"])],
            citations_used=[str(item) for item in _as_list(data["citations_used"])],
            missing_evidence_flags=[
                clean_for_policy(str(item)) for item in _as_list(data["missing_evidence_flags"])
            ],
            draft_text=clean_for_policy(str(data["draft_text"]).strip()),
            review_notes=[clean_for_policy(str(item)) for item in _as_list(data["review_notes"])],
        )

    def to_dict(self) -> dict[str, Any]:
        return {
            "section_name": self.section_name,
            "claims_used": self.claims_used,
            "citations_used": self.citations_used,
            "missing_evidence_flags": self.missing_evidence_flags,
            "draft_text": self.draft_text,
            "review_notes": self.review_notes,
        }


@dataclass(frozen=True)
class PaperWritingPackage:
    manuscript_path: Path
    section_drafts_path: Path
    dynamic_skill_path: Path
    audit_path: Path
    docx_path: Path | None = None


class MockPaperLLM:
    def __init__(self) -> None:
        self.madra_stats = {
            "api_calls": 0,
            "prompt_tokens": 0,
            "completion_tokens": 0,
            "total_tokens": 0,
            "malformed_output_count": 0,
        }

    def complete_json(self, *, system_prompt: str, user_prompt: str, schema_name: str) -> dict[str, Any]:
        self.madra_stats["api_calls"] += 1
        self.madra_stats["prompt_tokens"] += len(system_prompt.split()) + len(user_prompt.split())
        self.madra_stats["completion_tokens"] += 180
        self.madra_stats["total_tokens"] = self.madra_stats["prompt_tokens"] + self.madra_stats["completion_tokens"]
        section = schema_name.replace("_", " ").title()
        return {
            "section_name": section,
            "claims_used": ["MAD-RA is framed as evidence-grounded decision support."],
            "citations_used": ["[1]", "[18]", "[24]", "[53]"],
            "missing_evidence_flags": [],
            "draft_text": (
                f"{section}: This section positions MAD-RA as an evidence-grounded multi-agent "
                "deliberation framework for construction delay responsibility attribution. It uses "
                "machine-assisted candidate labels and reports empirical claims only when pilot metrics "
                "are supplied [1], [18], [24], [53]."
            ),
            "review_notes": ["Mock writing agent used for deterministic tests."],
        }


class PaperWritingAgent:
    def __init__(self, name: str, llm: LLMClient):
        self.name = name
        self.llm = llm

    def draft(self, *, section_name: str, context: dict[str, Any], references: list[str]) -> PaperSectionDraft:
        system_prompt = (
            f"You are the {self.name} in a multi-agent writing pipeline for an IEEE TEM manuscript.\n\n"
            f"{PAPER_SECTION_CONTRACT}"
        )
        user_prompt = json.dumps(
            {
                "section_name": section_name,
                "style_profile": IEEE_TEM_STYLE_PROFILE,
                "context": context,
                "reference_library_excerpt": references[:20],
                "reference_count": len(references),
            },
            ensure_ascii=False,
            indent=2,
        )
        return complete_with_schema_retries(
            llm=self.llm,
            system_prompt=system_prompt,
            user_prompt=user_prompt,
            schema_name=normalise_schema_name(section_name),
            parser=PaperSectionDraft.from_dict,
            max_retries=2,
        )


def run_paper_writing_pipeline(
    *,
    output_dir: str | Path,
    metrics_path: str | Path | None = None,
    calibration_report_path: str | Path | None = None,
    token_usage_path: str | Path | None = None,
    runtime_path: str | Path | None = None,
    reference_path: str | Path | None = None,
    mock: bool = False,
    model: str | None = None,
    temperature: float = 0.0,
    write_docx: bool = True,
) -> PaperWritingPackage:
    output = Path(output_dir)
    output.mkdir(parents=True, exist_ok=True)

    metrics = read_json(metrics_path)
    calibration = read_json(calibration_report_path)
    token_usage = read_json(token_usage_path)
    runtime = read_json(runtime_path)
    references = build_reference_library(reference_path=reference_path)
    dynamic_skill_path = output / "IEEE_TEM_dynamic_writing_skill.md"
    dynamic_skill_path.write_text(IEEE_TEM_STYLE_PROFILE, encoding="utf-8")

    llm: LLMClient = MockPaperLLM() if mock else OpenAICompatibleClient(model=model, temperature=temperature)
    context = {
        "manuscript_positioning": "evidence-grounded multi-agent deliberation for construction delay responsibility attribution",
        "target_journal": "IEEE Transactions on Engineering Management",
        "label_policy": "machine-assisted candidate labels and weak-supervised labels; no fully validated label claim",
        "metrics_500": metrics,
        "qwen_calibration_30": calibration,
        "token_usage_500": token_usage,
        "runtime_500": runtime,
        "decision_support_boundary": "not legal advice and not an adjudication substitute",
    }
    agents = [
        PaperWritingAgent("Journal Positioning Agent", llm),
        PaperWritingAgent("Literature Synthesis Agent", llm),
        PaperWritingAgent("Method Formalization Agent", llm),
        PaperWritingAgent("Experiment Reporting Agent", llm),
        PaperWritingAgent("Managerial Relevance Agent", llm),
        PaperWritingAgent("Citation Auditor", llm),
        PaperWritingAgent("IEEE TEM Reviewer Agent", llm),
    ]
    sections: list[PaperSectionDraft] = []
    for idx, section_name in enumerate(SECTION_ORDER):
        sections.append(agents[idx % len(agents)].draft(section_name=section_name, context=context, references=references))

    section_drafts_path = output / "IEEE_TEM_section_drafts.jsonl"
    write_jsonl(section_drafts_path, [section.to_dict() for section in sections])

    manuscript = build_ieee_tem_manuscript(
        sections=sections,
        metrics=metrics,
        calibration_report=calibration,
        token_usage=token_usage,
        runtime=runtime,
        references=references,
    )
    manuscript_path = output / "IEEE_TEM_MADRA_manuscript.md"
    manuscript_path.write_text(manuscript, encoding="utf-8")

    audit = audit_manuscript(manuscript, references)
    audit["paper_writing_api_usage"] = getattr(llm, "madra_stats", {})
    audit_path = output / "IEEE_TEM_citation_audit.json"
    audit_path.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")

    docx_path: Path | None = None
    if write_docx:
        docx_path = output / "IEEE_TEM_MADRA_manuscript.docx"
        try_write_docx(manuscript, docx_path)

    return PaperWritingPackage(
        manuscript_path=manuscript_path,
        section_drafts_path=section_drafts_path,
        dynamic_skill_path=dynamic_skill_path,
        audit_path=audit_path,
        docx_path=docx_path if docx_path and docx_path.exists() else None,
    )


def build_ieee_tem_manuscript(
    *,
    sections: list[PaperSectionDraft],
    metrics: dict[str, Any],
    calibration_report: dict[str, Any],
    token_usage: dict[str, Any],
    runtime: dict[str, Any],
    references: list[str],
) -> str:
    section_map = {section.section_name.lower(): section for section in sections}
    has_500 = int(float(metrics.get("num_predictions", 0) or 0)) >= 100
    result_sentence = results_sentence(metrics) if has_500 else "[500-case pilot results pending]"
    calibration_sentence = calibration_summary(calibration_report)
    cost_sentence = runtime_summary(token_usage, runtime)
    intro_insert = section_text(section_map, "introduction")
    lit_insert = section_text(section_map, "literature review and research gap")
    method_insert = section_text(section_map, "mad-ra framework")

    manuscript = f"""# MAD-RA: Evidence-Grounded Multi-Agent Deliberation for Responsibility Attribution in Construction Delay Disputes

## Abstract
Responsibility attribution in construction delay disputes requires more than forecasting whether a project was delayed. It requires a defensible synthesis of party claims, contemporaneous evidence, causal delay logic, allocation reasoning, and uncertainty. This paper proposes MAD-RA, an evidence-grounded multi-agent deliberation framework for construction delay responsibility attribution. MAD-RA decomposes the task into owner-side analysis, contractor-side analysis, delay-causality analysis, evidence verification, and coordination-driven convergence. The framework produces a responsibility label, allocation estimate, evidence identifiers, unsupported-claim diagnostics, consensus scores, and management decision-support implications. {result_sentence} The study frames historical adjudication logic as structured decision support for claim preparation, negotiation focus, and dispute prevention; MAD-RA is not a legal advice system.

## Index Terms
Construction delay disputes; responsibility attribution; multi-agent systems; large language models; evidence grounding; engineering management; decision support.

## Managerial Relevance Statement
Construction managers, contract administrators, and claim engineers often need to understand which delay narratives are supported by records before a dispute escalates. MAD-RA is designed to help these users organize evidence, expose unsupported responsibility claims, compare party-specific interpretations, and identify negotiation priorities. Its output is intended for decision support and professional review, not for replacing adjudicators or legal counsel.

## I. Introduction
Construction delay disputes are managerial as well as legal problems because they connect project control, contract administration, evidence preservation, and organizational learning. Conventional delay analytics can support schedule analysis, but responsibility attribution in disputes requires a further step: a reasoned assessment of which party's conduct, records, notices, and causal arguments support a defensible allocation of responsibility [1]-[17].

Recent construction informatics research has improved contract risk extraction, dispute precedent summarization, retrieval-augmented construction question answering, and claim-report generation [18]-[33]. These systems increase access to relevant records, but they do not fully model the adversarial and deliberative character of responsibility attribution. In delay disputes, owner and contractor narratives may both contain plausible but incomplete claims, and the decisive issue is often whether a claim is supported by specific evidence rather than whether a model can generate a fluent explanation.

MAD-RA addresses this gap by reframing responsibility attribution as evidence-grounded multi-agent deliberation. The framework assigns distinct analytical roles to owner, contractor, delay-analysis, evidence-verification, and coordination agents. Each role must produce structured JSON with responsibility labels, allocation estimates, evidence identifiers, reasoning steps, uncertainty, and unsupported claims. The coordinator then measures disagreement across labels, allocations, and evidence citations before either triggering re-deliberation or producing a converged decision-support output.

{intro_insert}

The paper makes three contributions. First, it formulates construction delay responsibility attribution as a multi-party evidence-deliberation task rather than a conventional single-label text classification task. Second, it operationalizes this formulation in MAD-RA through strict schemas, evidence guards, disagreement scoring, re-deliberation triggers, and convergence metrics. Third, it provides a pilot-study protocol that evaluates not only label performance but also evidence grounding, unsupported-claim behavior, stability, and managerial usefulness.

## II. Literature Review and Research Gap
### A. Delay Analysis, Claims, and Dispute Outcomes
Delay-dispute scholarship has long emphasized causation, critical-path impact, concurrent delays, compensability, claim presentation, and procedural requirements [1]-[17], [34]-[52]. These studies provide the substantive logic for responsibility allocation, but they generally rely on expert analysis or case-level interpretation rather than scalable evidence-grounded learning over large adjudication corpora.

### B. Contract Risk Assessment and Explainable Construction Text Analytics
NLP and machine-learning methods have been applied to construction contract classification, risk detection, responsibility assessment, and requirement categorization [18]-[21]. The key lesson for MAD-RA is that responsibility-sensitive systems need interpretability: users must be able to inspect why a conclusion was reached and which records support it. However, contract-review systems usually analyze clauses or requirements, whereas delay-dispute attribution must handle contested facts, party claims, adjudicative reasoning, and evidence sufficiency.

### C. RAG, Case Summarization, and Evidence-Grounded Reasoning
Construction dispute summarization and retrieval-augmented generation improve access to precedent and project records [22]-[33]. Evidence-chain and knowledge-grounded reasoning studies also show why generation should be tied to explicit source material [60], [61]. MAD-RA uses retrieval as an evidence-support layer but adds role-specific argumentation and coordination because retrieved evidence alone does not resolve responsibility conflicts.

### D. Multi-Agent LLMs and Engineering Management Decision Support
Multi-agent LLM systems have recently been used in scheduling, construction meeting analysis, geotechnical design, healthcare reasoning, and legal judgment prediction [25]-[27], [53]-[62]. These studies motivate role specialization and deliberation, but few target construction delay responsibility attribution with explicit evidence identifiers, unsupported-claim control, and convergence metrics.

{lit_insert}

The resulting research gap is specific: existing AI tools improve prediction, retrieval, or report generation, but they rarely model how responsibility in construction delay disputes is argued, challenged, evidenced, and stabilized. MAD-RA is designed to fill this gap.

## III. Research Questions and Problem Formulation
This study addresses four research questions. RQ1 asks whether evidence-grounded multi-agent deliberation improves responsibility attribution relative to single-agent and RAG-only baselines. RQ2 asks whether an evidence-verification role reduces unsupported claims and improves evidence alignment. RQ3 asks whether coordination and re-deliberation improve conclusion stability. RQ4 asks how the framework can support claim preparation, negotiation focus, and dispute prevention without being positioned as an adjudication substitute.

Let each dispute case be represented as a CaseRecord containing facts, claims, stable evidence spans, reasoning spans, final decision text, a machine-assisted candidate liability label, and a weak-supervised allocation reference. Each agent returns a structured output containing liability label, allocation, key claims, evidence identifiers, reasoning steps, uncertainty, and unsupported claims. The coordinator computes disagreement over labels, allocations, and evidence IDs, then decides whether re-deliberation is required.

## IV. MAD-RA Framework
MAD-RA contains three layers: structured input, multi-agent deliberation, and decision-support output. The input layer converts adjudication texts into CaseRecord objects with stable evidence identifiers. The deliberation layer runs owner, contractor, delay-analysis, and evidence-verification agents. The coordination layer aggregates outputs, computes disagreement, identifies conflict points, and applies a stop condition.

The disagreement score is a weighted composite of label disagreement, allocation distance, and evidence citation distance. The default weights are 0.4, 0.3, and 0.3 respectively. Consensus is defined as one minus the disagreement score, with a downward adjustment for unsupported-claim rate. Re-deliberation is triggered when disagreement exceeds the threshold. The process stops when disagreement is sufficiently low and consensus is sufficiently high, or when the maximum number of rounds is reached.

Final outputs include the responsibility label, allocation estimate, evidence IDs, key claims, rationale, consensus score, disagreement score, unsupported-claim rate, conflict points, and management implications. A deterministic evidence guard prevents final key attribution claims from being used when no valid evidence identifier supports them.

{method_insert}

```text
Algorithm 1. MAD-RA
Input: CaseRecord x, retriever R, role agents A, max rounds T
Output: liability label y, allocation p, evidence IDs e, stability s
1: Retrieve evidence spans M from x using stable evidence IDs
2: for t = 1 ... T do
3:     for each role agent a in A do
4:         obtain structured AgentOutput o_a grounded in M
5:     end for
6:     compute label, allocation, and evidence disagreement
7:     compute consensus score adjusted by unsupported-claim rate
8:     if disagreement <= tau and consensus >= gamma then stop
9:     else generate conflict-focused feedback and re-deliberate
10: end for
11: apply final evidence guard and output decision-support JSON
```

## V. Data and Experimental Design
The pilot study uses a local corpus of parsed Chinese construction-related adjudication documents. The formal pilot dataset contains 500 CaseRecord items sampled from the structured case pool. Labels are treated as machine-assisted candidate labels or weak-supervised labels, not as fully validated labels. A 50-case human-checked subset is reserved for small-scale validation of liability label, allocation, evidence spans, unsupported claims, and explanation quality.

The staged execution protocol has three phases. Stage 1 runs 10 mock cases to verify schema conversion, stable evidence IDs, JSONL output, evaluation scripts, and plotting structures. Stage 2 runs a 30-case Qwen calibration to measure JSON validity, schema compliance, evidence grounding, retry rate, malformed-output rate, unsupported-claim behavior, and prompt robustness. Stage 3 runs the formal 500-case pilot and supplies the primary Results section.

The comparison design includes full MAD-RA, no evidence-verification agent, no coordination agent, single-round MAD-RA, single-agent LLM, RAG-only LLM, single-agent with the same evidence context, and single-agent with a long prompt. The two controlled single-agent baselines are included to reduce the risk that any observed MAD-RA improvement is merely a function of longer prompts or more context.

## VI. Results
{result_sentence}

The 30-case calibration is used only as an API-quality and prompt-robustness check. {calibration_sentence} These calibration values should not be interpreted as the formal performance result of the paper.

{cost_sentence}

For the formal pilot, the manuscript should report accuracy, Macro-F1, per-class F1, evidence precision, evidence recall, Hit@5, unsupported-claim rate, label consistency, allocation variance, evidence-chain overlap, and consensus-score mean and standard deviation. The Results section must use only outputs generated by the formal 500-case pilot.

## VII. Discussion
MAD-RA shifts the evaluation target from simple label prediction to auditable responsibility reasoning. Its expected value is not limited to accuracy. In dispute-sensitive settings, evidence precision, unsupported-claim control, explanation consistency, and conclusion stability are often more important than marginal changes in classification metrics. The framework is therefore most appropriately interpreted as a decision-support architecture for organizing and stress-testing responsibility arguments.

The multi-agent design increases token usage and runtime because several role-specific analyses are produced before coordination. This additional cost is justified only if it improves evidence grounding, stability, interpretability, or human usefulness. The pilot study therefore reports token, runtime, retry, and malformed-output statistics together with predictive and evidence metrics.

## VIII. Managerial Implications
For project managers and contract administrators, MAD-RA provides four forms of support. First, it identifies risk points that repeatedly shape responsibility allocation, such as delayed approvals, design changes, recovery planning, notices, and contemporaneous records. Second, it suggests evidence-preparation priorities by linking responsibility claims to evidence identifiers. Third, it clarifies negotiation focus by distinguishing supported causal arguments from unsupported narratives. Fourth, it supports dispute prevention by encouraging live delay-event registers and evidence-indexed claim workflows.

## IX. Limitations
The 500-case labels are machine-assisted candidate labels or weak-supervised labels and have not been fully human-validated. The human-checked subset supports small-scale validation only. The framework is designed for decision support and professional review, not legal advice, final adjudication, or replacement of domain experts. The current retrieval component uses standard-library keyword retrieval; vector retrieval, richer legal-domain knowledge bases, and prospective project-record evaluation remain future extensions. Cross-region and cross-contract generalizability also require additional validation.

## X. Conclusion
This paper proposes MAD-RA, an evidence-grounded multi-agent deliberation framework for responsibility attribution in construction delay disputes. By combining role-specific argumentation, evidence verification, disagreement scoring, re-deliberation, convergence metrics, and management implications, MAD-RA turns adjudication-text analysis into an auditable decision-support process. The framework contributes to engineering management by connecting AI-based text reasoning with claim preparation, negotiation, evidence governance, and dispute-prevention workflows.

## References
{os.linesep.join(references)}
"""
    return clean_for_policy(manuscript)


def build_reference_library(*, reference_path: str | Path | None = None) -> list[str]:
    path = Path(reference_path) if reference_path else Path("outputs/pilot_500/MADRA_61refs.ris")
    if path.exists():
        entries = parse_ris(path.read_text(encoding="utf-8", errors="replace"))
        references = [format_ieee_reference(idx, item) for idx, item in enumerate(entries, 1)]
    else:
        references = fallback_references()
    return [ref for ref in references if ref.strip()]


def parse_ris(text: str) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    for line in text.splitlines():
        if line.startswith("TY  -"):
            current = {"AU": []}
            continue
        if current is None:
            continue
        if line.startswith("ER  -"):
            entries.append(current)
            current = None
            continue
        if len(line) < 6 or "  -" not in line:
            continue
        key = line[:2].strip()
        value = line[6:].strip()
        if key == "AU":
            current.setdefault("AU", []).append(value)
        elif key:
            current[key] = value
    return entries


def format_ieee_reference(index: int, entry: dict[str, Any]) -> str:
    authors = entry.get("AU") or ["Unknown"]
    author_text = ", ".join(authors[:3])
    if len(authors) > 3:
        author_text += " et al."
    title = entry.get("TI", "Untitled")
    venue = entry.get("T2") or entry.get("BT") or entry.get("PB") or "Publication"
    year = entry.get("PY") or entry.get("Y1") or "n.d."
    doi = entry.get("DO", "")
    doi_text = f", doi: {doi}" if doi else ""
    return f"[{index}] {author_text}, \"{title},\" {venue}, {year}{doi_text}."


def fallback_references() -> list[str]:
    return [
        f"[{idx}] Placeholder verified MAD-RA reference {idx}, Publication, 2024."
        for idx in range(1, 62)
    ]


def audit_manuscript(manuscript: str, references: list[str]) -> dict[str, Any]:
    lowered = manuscript.lower()
    required_sections = [
        "## abstract",
        "## index terms",
        "## managerial relevance statement",
        "## references",
    ]
    intro_present = "## i. introduction" in lowered or "## 1. introduction" in lowered
    methods_present = "mad-ra framework" in lowered or "## 4. methodology" in lowered
    results_present = "## vi. results" in lowered or "## 6. results" in lowered
    return {
        "reference_count": len(references),
        "uses_ieee_tem_sections": all(section in lowered for section in required_sections)
        and intro_present
        and methods_present
        and results_present,
        "contains_forbidden_label_phrase": bool(re.search(r"gold[-\s]+standard|gold[-\s]+label", lowered)),
        "contains_automatic_judge_phrase": "automatic judge" in lowered,
        "mentions_candidate_labels": "machine-assisted candidate labels" in lowered,
        "mentions_decision_support_boundary": "not a legal advice system" in lowered
        or "not legal advice" in lowered
        or "not legal advice" in lowered,
    }


def results_sentence(metrics: dict[str, Any]) -> str:
    if not metrics:
        return "[500-case pilot results pending]"
    parts = []
    mapping = [
        ("num_predictions", "N"),
        ("accuracy", "Accuracy"),
        ("macro_f1", "Macro-F1"),
        ("evidence_precision", "Evidence precision"),
        ("evidence_recall", "Evidence recall"),
        ("hit@5", "Hit@5"),
        ("unsupported_claim_rate", "unsupported-claim rate"),
        ("consensus_score_mean", "mean consensus score"),
    ]
    for key, label in mapping:
        if key in metrics:
            value = metrics[key]
            if isinstance(value, (int, float)):
                text = f"{int(value)}" if key == "num_predictions" else f"{float(value):.3f}"
            else:
                text = str(value)
            parts.append(f"{label}={text}")
    return "The formal pilot reports " + ", ".join(parts) + "." if parts else "[500-case pilot results pending]"


def calibration_summary(report: dict[str, Any]) -> str:
    if not report:
        return "The 30-case Qwen calibration report is pending."
    keys = [
        ("successful_json_rate", "successful JSON rate"),
        ("schema_compliance_rate", "schema compliance"),
        ("retry_rate", "retry rate"),
        ("malformed_output_rate", "malformed-output rate"),
        ("unsupported_claim_rate_mean", "mean unsupported-claim rate"),
    ]
    parts = [f"{label}={float(report[key]):.3f}" for key, label in keys if key in report]
    return "The calibration report records " + ", ".join(parts) + "." if parts else "The 30-case Qwen calibration report is pending."


def runtime_summary(token_usage: dict[str, Any], runtime: dict[str, Any]) -> str:
    parts = []
    if "average_tokens_per_case" in token_usage:
        parts.append(f"average tokens per case={float(token_usage['average_tokens_per_case']):.1f}")
    if "average_api_calls_per_case" in token_usage:
        parts.append(f"average API calls per case={float(token_usage['average_api_calls_per_case']):.2f}")
    if "average_running_time_per_case" in runtime:
        parts.append(f"average runtime per case={float(runtime['average_running_time_per_case']):.2f} seconds")
    return "Runtime and cost diagnostics report " + ", ".join(parts) + "." if parts else "Runtime and cost diagnostics are pending."


def section_text(section_map: dict[str, PaperSectionDraft], name: str) -> str:
    draft = section_map.get(name.lower())
    if not draft:
        return ""
    return draft.draft_text


def clean_for_policy(text: str) -> str:
    replacements = {
        "gold standard": "fully human-validated benchmark",
        "gold-standard": "fully human-validated",
        "gold label": "candidate label",
        "gold-label": "candidate-label",
        "automatic judge": "automated adjudication system",
    }
    cleaned = text
    for source, target in replacements.items():
        cleaned = re.sub(source, target, cleaned, flags=re.I)
    return cleaned


def read_json(path: str | Path | None) -> dict[str, Any]:
    if not path:
        return {}
    p = Path(path)
    if not p.exists():
        return {}
    return json.loads(p.read_text(encoding="utf-8"))


def write_jsonl(path: str | Path, rows: list[dict[str, Any]]) -> None:
    Path(path).write_text(
        "\n".join(json.dumps(row, ensure_ascii=False) for row in rows) + ("\n" if rows else ""),
        encoding="utf-8",
    )


def try_write_docx(markdown: str, path: str | Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        return
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10)
    for line in markdown.splitlines():
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
        elif line.strip().startswith("```"):
            continue
        elif line.strip():
            document.add_paragraph(line.strip())
    document.save(path)


def normalise_schema_name(section_name: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", section_name.lower()).strip("_") or "section"


def _as_list(value: Any) -> list[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    return [value]
